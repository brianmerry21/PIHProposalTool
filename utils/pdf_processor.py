import pdfplumber
# import pandas as pd  # Temporarily commented out due to NumPy dependency issue
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
import logging
import os
import re
from datetime import datetime
import tempfile
from pypdf import PdfReader
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from xml.etree.ElementTree import fromstring as parse_xml
from utils.image_processor import extract_pdf_region_as_image, extract_page_as_image_special, process_logo_image


from pdf2image import convert_from_path
from docx.enum.section import WD_SECTION
# Set up logging
logger = logging.getLogger(__name__)

VALUE = 'software'

def set_image_in_front_of_text(run):
    """
    Set the image in a run to be positioned in front of text
    (allows for larger images that can extend into margins)

    Args:
        run: The run containing the image
    """
    try:
        # Get the drawing element (only exists if the run contains an image)
        drawing_element = None
        for child in run._element:
            if child.tag.endswith(('drawing')):
                drawing_element = child
                break

        if drawing_element is not None:
            # Find the appropriate elements to modify
            inline = drawing_element.find('.//wp:inline', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            if inline is not None:
                # Change inline to anchor to enable text wrapping
                anchor = OxmlElement('wp:anchor')
                anchor.set('distT', '0')
                anchor.set('distB', '0')
                anchor.set('distL', '0')
                anchor.set('distR', '0')
                anchor.set('simplePos', '0')
                anchor.set('relativeHeight', '1')
                anchor.set('behindDoc', '0')
                anchor.set('locked', '0')
                anchor.set('layoutInCell', '1')
                anchor.set('allowOverlap', '1')

                # Copy attributes and children from inline to anchor
                for key, value in inline.attrib.items():
                    if key != 'distT' and key != 'distB' and key != 'distL' and key != 'distR':
                        anchor.set(key, value)

                for child in inline:
                    anchor.append(child)

                # Add required child elements for anchor
                simple_pos = OxmlElement('wp:simplePos')
                simple_pos.set('x', '0')
                simple_pos.set('y', '0')
                anchor.insert(0, simple_pos)

                # Position relative to page (not margin)
                pos_h = OxmlElement('wp:positionH')
                pos_h.set('relativeFrom', 'page')
                pos_h_align = OxmlElement('wp:align')
                pos_h_align.text = 'center'
                pos_h.append(pos_h_align)
                anchor.append(pos_h)

                pos_v = OxmlElement('wp:positionV')
                pos_v.set('relativeFrom', 'page')
                pos_v_offset = OxmlElement('wp:posOffset')
                # Position slightly down from the top of the page to avoid headers
                pos_v_offset.text = '1250000'  # EMUs (English Metric Units)
                pos_v.append(pos_v_offset)
                anchor.append(pos_v)

                # Set text wrapping to "in front of text"
                wrap_none = OxmlElement('wp:wrapNone')
                anchor.append(wrap_none)

                # Replace the inline element with our new anchor element
                drawing_element.remove(inline)
                drawing_element.append(anchor)

                logger.debug("Successfully set image to 'In Front of Text'")
                return True

        return False
    except Exception as e:
        logger.warning(f"Error setting image wrapping: {e}")
        return False

# new function to extract from the cost sheet to doc

import openpyxl

def extract_optional_items_from_excel(excel_path):
    """Extract optional items from Excel (B37:B75=description, C37:C75=price estimation, K37:K75=quantity)."""
    logger.info(f"Starting Excel extraction from: {excel_path}")
    
    try:
        global VALUE 
        
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        ws = wb.active  
        logger.info(f"Excel file loaded successfully. Active sheet: {ws.title}")
       
        # B18 -> column B, row 18
        VALUE = ws["B18"].value
        if isinstance(VALUE, str):
            VALUE = VALUE.strip()
            
        print(f"VALUE is {VALUE}")    


        # Debug log a wider range
        logger.info("Checking Excel file content...")
        for row in range(30, 80):
            b_val = ws[f"B{row}"].value
            c_val = ws[f"C{row}"].value
            k_val = ws[f"K{row}"].value
            if b_val or c_val or k_val:
                logger.info(f"Row {row}: B={b_val}, C={c_val}, K={k_val}")

        items = []
        for row in range(37, 76):  # B37:B75, C37:C75, K37:K75
            description_val = ws[f"B{row}"].value  # description
            price_val = ws[f"C{row}"].value       # price estimation
            quantity_val = ws[f"K{row}"].value    # quantity

            # Skip if row is blank
            if not any([description_val, price_val, quantity_val]):
                continue
            if not description_val or str(description_val).strip() == "":
                continue

            description = str(description_val).strip()

            # Handle price estimation
            price_each = 0
            is_tbd_price = False
            if price_val is None or str(price_val).strip() == "":
                is_tbd_price = True
            else:
                price_str = str(price_val).strip()
                if price_str.upper() == "TBD":
                    is_tbd_price = True
                else:
                    try:
                        # Remove $ and commas
                        cleaned = price_str.replace("$", "").replace(",", "")
                        price_each = float(cleaned)
                    except Exception:
                        logger.warning(f"Invalid price in row {row}: {price_val}, marking as TBD")
                        is_tbd_price = True
                        price_each = 0


            # Handle quantity
            try:
                quantity = int(quantity_val) if quantity_val not in (None, "") else 1
            except Exception:
                logger.warning(f"Invalid quantity in row {row}: {quantity_val}, defaulting to 1")
                quantity = 1

            items.append({
                "description": description,
                "price_each": price_each,
                "quantity": quantity,
                "is_tbd_price": is_tbd_price,
            })

        logger.info(f"Total items extracted: {len(items)}")
        return items

    except Exception as e:
        logger.error(f"Error in extract_optional_items_from_excel: {str(e)}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return []


#end of the function


#new brian code 
from openpyxl import load_workbook

NBSP = "\u00A0"

def _normalize(s: str) -> str:
    return " ".join(s.replace(NBSP, " ").split()).strip()

def _value_from_cell_or_merge(ws, addr: str):
    c = ws[addr]
    if c.value is not None:
        return c.value
    # If B18 is inside a merged range, the value lives at the top-left cell of that range
    for mr in ws.merged_cells.ranges:
        if c.coordinate in mr:
            return ws.cell(row=mr.min_row, column=mr.min_col).value
    return None

def get_b18_value(excel_path):
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active  # Use the first sheet; change if needed

    # B18 -> column B, row 18
    value = ws["B18"].value
    if isinstance(value, str):
        value = value.strip()
    return value


#new brian code


def extract_customer_info_from_pdf(pdf_text):
    """

    Args:

    Returns:
        Dictionary with customer info (name, location, date, proposal_number, contact_info)
    """
    import re
    import logging
    from datetime import datetime
    logger = logging.getLogger(__name__)

    customer_info = {
        "name": "",  # Default
        "location": "Hillsboro, OR",  # Default
        "date": datetime.now().strftime('%B %d, %Y'),
        "proposal_number": "00000000",
        "model": "Modula ML25 Vertical Lift Module",
        "contact_name": "Josh Jancola",
        "contact_email": "joshjancola@pacificintegrated.com",
        "contact_phone": "253.500.4193",
        "contact_office": "888.550.5888"
    }

    # Try to extract customer name from the header
    import re

    # Extract the first few hundred characters from the PDF text
    text_snippet = full_text[:1000]  # increase range if necessary


    # Get only the first page text
    first_page_text = pdf_text_pages[0]

    # (Optional) Remove known headers/footers
    first_page_text = first_page_text.replace("Dart Aerospace", "")

    # Regex for possible customer names (avoid plain word "Customer")
    pattern = re.compile(
        r"\b([A-Z][A-Za-z&,\.\s]*(?:Inc\.?|Corp\.?|Corporation|Company|LLC|Scientific)?)\b"
    )

    # Search for match
    customer_match = pattern.search(first_page_text)

    if customer_match:
        client_name = customer_match.group(1).strip()

        # Avoid picking "Customer" as the name
        if client_name.lower() == "customer":
            client_name = "Unknown Customer"
            print("Matched 'Customer' instead of real name.")
        else:
            print(f"Found customer name: {client_name}")
            print("Extracted from snippet:", first_page_text[
                customer_match.start()-30 : customer_match.end()+30
            ])
    else:
        client_name = "Unknown Customer"
        print("No customer name match found.")

    # Try to extract proposal number (based on screenshot format)
    proposal_pattern = r"Proposal\s+[#]?([0-9]+)"
    proposal_match = re.search(proposal_pattern, pdf_text[:1000])
    if proposal_match:
        customer_info["proposal_number"] = proposal_match.group(1)

    # Try to extract location - look for cleanroom location
    location_pattern = r"(?:cleanroom at the|facility at|facility in)\s+(.*?)(?:facility|\.)"
    location_match = re.search(location_pattern, pdf_text[:2000])
    if location_match:
        location = location_match.group(1).strip()
        if location:
            customer_info["location"] = location

    # Backup location pattern
    if "location" not in customer_info or not customer_info["location"]:
        location_backup = r"(?:Hillsboro|Seattle|Portland|Tacoma|Bellevue|Vancouver),\s+(?:OR|WA|CA|ID)"
        location_match = re.search(location_backup, pdf_text[:2000])
        if location_match:
            customer_info["location"] = location_match.group(0)

    # Try to extract date from the header
    date_pattern = r"((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4})"
    date_match = re.search(date_pattern, pdf_text[:1000])
    if date_match:
        customer_info["date"] = date_match.group(1)

    # Try to extract model from the header or content
    model_pattern = r"Modula\s+(?:VLM\s+)?(?:ML|MC|MX|MS)\d+(?:-\d+)?"
    model_match = re.search(model_pattern, pdf_text[:2000])
    if model_match:
        customer_info["model"] = model_match.group(0)

    # Try to extract tray information
    tray_pattern = r"contain\s+(\d+)\s+slotted\s+trays"
    tray_match = re.search(tray_pattern, pdf_text)
    if tray_match:
        customer_info["num_trays"] = tray_match.group(1)

    # Try to extract tray dimensions
    dimension_pattern = r"sized\s+as\s+([\d\.]+)(?:\"|\s+inches|\s*\')\s+wide\s+and\s+([\d\.]+)(?:\"|\s+inches|\s*\')\s+deep"
    dimension_match = re.search(dimension_pattern, pdf_text)
    if dimension_match:
        width = dimension_match.group(1)
        depth = dimension_match.group(2)
        customer_info["tray_dimensions"] = f"Width: {width}'' x Depth: {depth}''"

    # Try to extract contact info (based on screenshot format)
    contact_name_pattern = r"([A-Z][a-z]+\s+[A-Z][a-z]+)[\s\n]*Pacific\s+Integrated\s+Handling"
    contact_name_match = re.search(contact_name_pattern, pdf_text)
    if contact_name_match:
        customer_info["contact_name"] = contact_name_match.group(1)

    # Try to extract phone
    phone_pattern = r"Cell:\s+([\d\.\-]+)"
    phone_match = re.search(phone_pattern, pdf_text)
    if phone_match:
        customer_info["contact_phone"] = phone_match.group(1)

    # Try to extract email
    email_pattern = r"\b([A-Za-z0-9._%+-]+@pacificintegrated\.com)\b"
    email_match = re.search(email_pattern, pdf_text)
    if email_match:
        customer_info["contact_email"] = email_match.group(1)

    # Look for recipient name (person being addressed)
    recipient_pattern = r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+),\s*[\n\r]"
    recipient_match = re.search(recipient_pattern, pdf_text[:2000])
    if recipient_match:
        customer_info["recipient_name"] = recipient_match.group(1)

    # Extract storage capacity information
    capacity_pattern = r"storage\s+capacity\s+of\s+([\d\.]+)\s+sq\s+ft\.\s+and\s+([\d\.]+)\s+cubic\s+ft\."
    capacity_match = re.search(capacity_pattern, pdf_text)
    if capacity_match:
        sq_ft = capacity_match.group(1)
        cubic_ft = capacity_match.group(2)
        customer_info["storage_capacity"] = f"{sq_ft} sq ft. / {cubic_ft} cubic ft."

    # Extract footprint
    footprint_pattern = r"footprint\s+of\s+([\d\.]+)\s+sq\.\s+ft\."
    footprint_match = re.search(footprint_pattern, pdf_text)
    if footprint_match:
        customer_info["unit_footprint"] = f"{footprint_match.group(1)} sq. ft."

    # Extract height
    height_pattern = r"height\s+of\s+([\d\.]+)\s+inches\s+\(([\d\'\.\"]+)\)"
    height_match = re.search(height_pattern, pdf_text)
    if height_match:
        inches = height_match.group(1)
        feet = height_match.group(2)
        customer_info["unit_height"] = f"{inches} inches ({feet})"
        
    return customer_info

def format_tray_dimensions(text):
    """
    Extract and format tray dimensions from PDF text

    Args:
        text: Full text content from PDF

    Returns:
        Formatted string with tray width and depth
    """
    try:
        # Extract width
        width_match = re.search(r"Usable tray width:\s+(\d+\.\d+)\s+in", text)
        width = width_match.group(1) if width_match else "N/A"

        # Extract depth
        depth_match = re.search(r"Usable tray depth:\s+(\d+\.\d+)\s+in", text)
        depth = depth_match.group(1) if depth_match else "N/A"

        # Format as required
        return f"Width: {width}'' x Depth: {depth}''"
    except Exception as e:
        logger.error(f"Error formatting tray dimensions: {str(e)}")
        return "Width: N/A'' x Depth: N/A''"

def extract_table_items(text):
    """
    Extract line items from the table on pages 6-7 of the PDF

    Args:
        text: Full text content from PDF

    Returns:
        List of item descriptions to be added in column B
    """
    try:
        # This will store all our extracted items
        items = []

        # We'll search for items in sections that typically contain the parts table
        # Look for lines that typically represent items in the table
        # Using a pattern that matches item descriptions in the Modula PDFs
        item_pattern = r"(?:^|\n)([A-Z0-9]+-[A-Z0-9]+\s+.*?)(?:\n|$)"
        part_pattern = r"(?:^|\n)((?:[A-Za-z]+-)?[A-Z0-9]+-[A-Z0-9]+.*?)(?:\s+\d|\n|$)"

        # Try different patterns to capture items
        matches = re.findall(part_pattern, text, re.MULTILINE)
        if matches:
            for match in matches:
                # Clean up the match and filter out any non-item lines
                item = match.strip()
                # Only include lines that look like part numbers (contain alphanumeric with dash)
                if re.search(r"[A-Z0-9]+-[A-Z0-9]+", item) and len(item) > 5:
                    # Remove any trailing numbers or prices
                    item = re.sub(r"\s+\d+\.\d+$", "", item)
                    items.append(item)

        # Also try to find items in sections marked by Modula section names
        section_pattern = r"(?:Machine Section|Software Section|Install Section|Optional Items)[^\n]*\n(.*?)(?:\n\s*\n|\n\s*Total|\Z)"
        section_matches = re.findall(section_pattern, text, re.DOTALL)

        for section in section_matches:
            lines = section.strip().split('\n')
            for line in lines:
                line = line.strip()
                # Skip empty lines and numeric-only lines
                if line and not re.match(r"^\s*\d+\s*$", line):
                    # Check if it has a part number format
                    if re.search(r"[A-Z0-9]+-[A-Z0-9]+", line):
                        # Clean up line
                        line = re.sub(r"\s+\d+\.\d+$", "", line)
                        if line not in items:
                            items.append(line)

        # If no items found, try a more general approach
        if not items:
            # Look for table-like structures with part number patterns
            table_pattern = r"(?:Item|Part|Description)[^\n]*(?:\n.*){2,}?"
            table_matches = re.findall(table_pattern, text, re.IGNORECASE | re.DOTALL)

            for table_section in table_matches:
                lines = table_section.strip().split('\n')
                for line in lines[1:]:  # Skip header row
                    line = line.strip()
                    if line and re.search(r"[A-Z0-9]+-[A-Z0-9]+", line):
                        # Remove quantities and prices
                        clean_line = re.sub(r"\s+\d+\s+[\d,]+\.\d+", "", line)
                        clean_line = re.sub(r"\s+\d+\.\d+$", "", clean_line)
                        if clean_line not in items:
                            items.append(clean_line)

        # Remove duplicates while maintaining order
        unique_items = []
        for item in items:
            if item not in unique_items:
                unique_items.append(item)

        return unique_items

    except Exception as e:
        logger.error(f"Error extracting table items: {str(e)}")
        return ["Error extracting table items"]

def extract_optional_items(text):
    """
    Extract optional items from the pricing tables in the PDF

    Args:
        text: Full text content from PDF

    Returns:
        List of dictionaries containing optional item details
    """
    optional_items = []

    try:
        # More comprehensive list of optional accessories/components based on screenshot
        optional_item_keywords = [
            "SLIDING OPERATOR CONSOLE",
            "BAR CODE READER",
            "WIRELESS BAR CODE READER",
            "RFID",
            "BADGE READER",
            "MAGNETIC BADGE READER",
            "SPARE PARTS KIT",
            "PROTECTION SYSTEM",
            "UPGRADE TO PREMIUM",
            "TELEPHONE SUPPORT",
            "MAGNETIC",
            "ESD",
            "ELECTROSTATIC DISCHARGE",
            "LEVEL A",
            "LEVEL B",
            "LEVEL C",
            "HEAVY DUTY",
            "CONSOLE",
            "TRAY PARTITIONS",
            "DIVIDERS",
            "PUT-TO-LIGHT",
            "PUT TO LIGHT",
            "SPARE PARTS",
            "1D",
            "2D",
            "24/5",
            "24/7",
            "PREMIUM"
        ]

        # Check for a separate "Optional Items" section
        optional_section_indicators = ["OPTIONAL ITEMS", "OPTIONS NOT INCLUDED", "ACCESSORIES PRICING"]

        # Find if there's a dedicated optional items section
        section_start_idx = -1
        for indicator in optional_section_indicators:
            section_idx = text.upper().find(indicator)
            if section_idx > -1:
                section_start_idx = section_idx
                break

        # If we found a section, extract all lines from that section
        if section_start_idx > -1:
            # Get the relevant portion of text containing optional items
            section_text = text[section_start_idx:]
            # Split into lines and process each line
            section_lines = section_text.split('\n')

            # Variables to track when we're inside the options table
            in_options_section = True

            for line in section_lines:
                line = line.strip()

                # Skip empty or very short lines
                if not line or len(line) < 5:
                    continue

                # Look for price patterns in the line
                price_match = re.search(r'\$\s*([\d,]+(?:\.\d+)?)', line)

                if price_match and any(keyword.lower() in line.lower() for keyword in optional_item_keywords):
                    # Extract the price
                    price_str = price_match.group(1).replace(',', '')
                    try:
                        price_value = float(price_str)
                    except ValueError:
                        price_value = 0.0

                    # Extract the description (text before the price)
                    description = line.split('$')[0].strip()

                    # Add to our list of optional items
                    optional_items.append({
                        "description": description,
                        "price_each": price_value,
                        "quantity": 1,
                        "is_optional": True,
                        "is_tbd_price": False,
                        "category": "Optional"
                    })

        # If we didn't find items in a dedicated section, look throughout the document
        if len(optional_items) == 0:
            lines = text.split('\n')
            for i, line in enumerate(lines):
                line = line.strip()

                # Skip short lines
                if not line or len(line) < 10:
                    continue

                # Check if this line contains any of our optional item keywords
                is_optional_item = False
                for keyword in optional_item_keywords:
                    if keyword in line.upper():
                        is_optional_item = True
                        break

                if is_optional_item and '$' in line:
                    # Try to extract price
                    price_match = re.search(r'\$\s*([\d,]+(?:\.\d+)?)', line)
                    price_value = 0.0
                    if price_match:
                        price_str = price_match.group(1).replace(',', '')
                        try:
                            price_value = float(price_str)
                        except ValueError:
                            pass

                    # Extract description (text before the price)
                    description = line.split('$')[0].strip()

                    # Default quantity is 1
                    quantity = 1

                    # Check if this item is already in our list
                    if not any(item["description"].lower() == description.lower() for item in optional_items):
                        optional_items.append({
                            "description": description,
                            "price_each": price_value,
                            "quantity": quantity,
                            "is_optional": True,
                            "is_tbd_price": False,
                            "category": "Optional"
                        })

    except Exception as e:
        logger.error(f"Error extracting optional items: {str(e)}")

    # Ensure we have these standard items
    standard_items = [
        {
            "description": "Tray Partitions and Dividers",
            "price_each": 0.0,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": True
        },
        {
            "description": "Put-to-light System",
            "price_each": 0.0,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": True
        },
        {
            "description": "Sliding Operator Console HEAVY DUTY",
            "price_each": 1535.22,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "Wireless Bar Code Reader (2D)",
            "price_each": 3323.10,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "ESD (Electrostatic Discharge) Protection System",
            "price_each": 1329.24,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "RFID Badge Reader",
            "price_each": 1121.55,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "Spare Parts Kit - Level A (basic)",
            "price_each": 1938.82,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "Spare Parts Kit - Level B (recommended A+B)",
            "price_each": 10490.47,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "Spare Parts Kit - Level C (Advanced A+B+C)",
            "price_each": 18654.60,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "Magnetic Badge Reader",
            "price_each": 671.68,
            "quantity": 1, 
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "Wireless Bar Code Reader (1D)",
            "price_each": 1985.55,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "Upgrade To Premium 24/5 (24 Hrs./Day, 5 Days/Week) Telephone Support Agreement (Year 1 & 2)",
            "price_each": 5000.00,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        },
        {
            "description": "Upgrade To Premium 24/7 (24 Hrs./Day, 7 Days/Week) Telephone Support Agreement (Years 1 & 2)",
            "price_each": 8250.00,
            "quantity": 1,
            "is_optional": True,
            "is_tbd_price": False
        }
    ]

    # Add standard items if they're not already present
    for std_item in standard_items:
        if not any(item["description"].lower() == std_item["description"].lower() for item in optional_items):
            optional_items.append({
                "description": std_item["description"],
                "price_each": std_item["price_each"],
                "quantity": std_item["quantity"],
                "is_optional": True,
                "is_tbd_price": std_item["is_tbd_price"],
                "category": "Optional"
            })

    return optional_items

def extract_total_value(text):
    """
    Extract the total dollar value from the PDF

    Args:
        text: Full text content from PDF

    Returns:
        The total value as a string with $ prefix
    """
    try:
        logger.info("Attempting to extract total dollar value")

        # First look for clear total pattern - "Total" followed by numbers
        total_patterns = [
            # Standard total format
            r"(?:Total|TOTAL)[^\n]*?[^\d](\d{1,3}(?:,\d{3})*\.\d{2})",
            # Total with newlines between
            r"(?:Total|TOTAL)(?:[^\n]*\n)*?[^\d](\d{1,3}(?:,\d{3})*\.\d{2})",
            # Total amount at the end of a line 
            r"(?:Total|TOTAL)[^\n]*?[\s:]?(\d{1,3}(?:,\d{3})*\.\d{2})",                 
            # Specific to Modula PDFs total format
            r"(?:Total\s+(?:Price|Amount)|TOTAL\s+(?:PRICE|AMOUNT))[^\n]*?[^\d](\d{1,3}(?:,\d{3})*\.\d{2})"
        ]

        for pattern in total_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                total_value = match.group(1)
                logger.info(f"Found total value: ${total_value}")
                return f"${total_value}"

        # Look for a section that might contain the total
        total_section_pattern = r"(?:Optional Items|Install Section|Total).*?(\d{1,3}(?:,\d{3})*\.\d{2})"
        section_match = re.search(total_section_pattern, text, re.IGNORECASE | re.DOTALL)
        if section_match:
            total_value = section_match.group(1)
            logger.info(f"Found total value from section: ${total_value}")
            return f"${total_value}"

        # If all else fails, try to find a price that looks like a total
        # Often the largest price near the end of the document
        all_amounts = re.findall(r"(\d{1,3}(?:,\d{3})*\.\d{2})", text)
        if all_amounts:
            # Try to get the largest amount, which is often the total
            try:
                # Convert to floats and find the largest
                float_amounts = [float(amt.replace(',', '')) for amt in all_amounts]
                largest_amount = all_amounts[float_amounts.index(max(float_amounts))]
                logger.info(f"Found largest amount: ${largest_amount}")
                return f"${largest_amount}"
            except:
                # Or just use the last amount as a fallback
                last_amount = all_amounts[-1]
                logger.info(f"Using last amount as total: ${last_amount}")
                return f"${last_amount}"

        logger.warning("Could not find total value in PDF")
        return "$0.00"  # Default if no match found

    except Exception as e:
        logger.error(f"Error extracting total value: {str(e)}")
        return "$0.00"

def extract_text_from_pdf(pdf_path):
    """
    Extract all text content from a PDF file

    Args:
        pdf_path: Path to the PDF file

    Returns:
        A list of strings, each representing a page of text
    """
    text_content = []

    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"Processing {total_pages} pages from PDF")

            for page_num, page in enumerate(pdf.pages):
                try:
                    logger.debug(f"Extracting text from page {page_num+1}/{total_pages}")
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                    else:
                        logger.warning(f"No text extracted from page {page_num+1}")
                        text_content.append(f"[No text content found on page {page_num+1}]")
                except Exception as page_e:
                    logger.warning(f"Error extracting text from page {page_num+1}: {str(page_e)}")
                    text_content.append(f"[Error extracting text from page {page_num+1}: {str(page_e)}]")
                    continue
    except Exception as e:
        logger.error(f"Error opening or processing PDF: {str(e)}")
        # Return error message as text instead of raising exception
        return [f"Error processing PDF: {str(e)}"]

    if not text_content:
        logger.warning("No text could be extracted from any page of the PDF")
        return ["No text content could be extracted from the PDF."]

    return text_content

# Using extract_pdf_region_as_image from utils.image_processor instead

def extract_tables_from_pdf(pdf_path):
    """
    Extract all tables from a PDF file

    Args:
        pdf_path: Path to the PDF file

    Returns:
        A list of pandas DataFrames, each representing a table
    """
    tables = []

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    # Try to extract tables from this page
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        if table:
                            # Convert to pandas DataFrame
                            # Make sure to handle empty column names
                            column_names = [f"Col_{i}" if col is None or col == "" else col 
                                           for i, col in enumerate(table[0])]

                            # Check if there's data after the header row
                            if len(table) > 1:
                                # Convert table data to basic list format
                                table_data = {
                                    'columns': column_names,
                                    'data': table[1:]
                                }
                                tables.append(table_data)
                except Exception as page_e:
                    # Log the error but continue processing other pages
                    logger.warning(f"Error extracting tables from page {page_num+1}: {str(page_e)}")
                    continue
    except Exception as e:
        logger.error(f"Error extracting tables from PDF: {str(e)}")
        # Don't raise the exception, just return any tables we managed to extract

    return tables

def process_pdf_to_excel(pdf_path, excel_path, template_path=None, customer_info=None):
    """
    Process a PDF file and save extracted data to an Excel file

    Args:
        pdf_path: Path to the PDF file
        excel_path: Path to save the Excel file
        template_path: Optional path to an Excel template file
        customer_info: Optional dictionary with customer information to override PDF extraction
    """
    try:
        # If a template is provided, use it to create the Excel file
        if template_path and os.path.exists(template_path):
            try:
                from utils.excel_template_analyzer import process_pdf_to_excel_with_template

                # Create customized mapping rules for the specific Modula PDF structure
                mapping_rules = create_modula_mapping_rules()

                # Process PDF and populate template
                process_pdf_to_excel_with_template(pdf_path, template_path, excel_path, mapping_rules, customer_info)
                logger.debug(f"Excel file created from template at {excel_path}")
                return
            except Exception as template_e:
                logger.error(f"Error using template method: {str(template_e)}")
                logger.info("Falling back to basic extraction method")
                # Continue with basic extraction if template method fails

        # Extract tables from the PDF
        logger.info("Using basic table extraction method")
        tables = extract_tables_from_pdf(pdf_path)

        # If no tables were found, try to extract text and convert to structured format
        if not tables:
            logger.info("No tables found, extracting text content")
            text_content = extract_text_from_pdf(pdf_path)
            # Create a simple DataFrame from text
            data = []
            for page_num, page_text in enumerate(text_content):
                lines = pagetext.split('\n')
                for line in lines:
                    data.append({
                        'Page': page_num + 1,
                        'Content': line.strip()
                    })

            # Create a simple data structure from the text content
            simple_data = {'data': data}
            tables = [simple_data]

        # Save tables to Excel file - temporarily disabled due to pandas dependency
        logger.info(f"Excel generation temporarily disabled - found {len(tables)} data structures")
        # Excel functionality will be restored once NumPy dependency is resolved

        logger.debug(f"Excel file created at {excel_path}")

    except Exception as e:
        logger.error(f"Error processing PDF to Excel: {str(e)}")

        # Excel generation temporarily disabled due to dependency issues
        logger.info("Excel generation temporarily disabled - error handling skipped")
        return


def create_modula_mapping_rules():
    """
    Create mapping rules specifically for Modula PDFs (first page extraction)
    Returns:
        Dictionary defining rules to map Modula PDF data to Excel cells
    """
    mapping_rules = {
        "Primary Option": {
            # Customer information
            "F1": {"default": "Dart Aerospace"},  # Hardcoded

            # Model information
            "B7": {
            "pattern": r"(?:MODULA[^\n]*Model\s+[A-Z0-9\-]+)",
            "group": 0,
            "transform": lambda x: x.strip()
            },

            # Total dollar value in cell G7 - exact value from screenshot
            "G7": {"default": ""},  # Hardcoded

            # Offer reference number
            "B8": {"pattern": r"Ref\. Offer No\.: QLI-(\d+)", "group": 1, "transform": lambda x: f"Offer # {x}"},

            # Number of trays
            "B9": {
                "pattern": r"Number of trays:\s*(\d+)",
                "group": 1,
                "transform": lambda x: f"Number of trays: {x} pcs"
            },

            # Tray dimensions (custom handler)
            "B10": {
                "custom_handler": lambda text: format_tray_dimensions(text)
            },

            # Height
            "B11": {
                "pattern": r"Height:\s*([\d.]+)\s*in",
                "group": 1,
                "transform": lambda x: f"Height: {x} in"
            },

            # Storage Volume
            "B12": {
                "pattern": r"Storage Volume:\s*([\d.]+)\s*ft³",
                "group": 1,
                "transform":lambda x: f"Storage Volume :{x} ft³"
            },

            # Storage Area
            "B13": {
                "pattern": r"Storage Area:\s*([\d.]+)\s*ft²",
                "group": 1,
                "transform": lambda x: f"Storage Area: {x} ft²"
            },

            # Unit Footprint
            "B14": {
                "pattern": r"Unit Footprint:\s*([\d.]+)\s*ft²",
                "group": 1,
                "transform": lambda x: f"Unit Footprint: {x} ft²"
            },
            "C37": {
                "pattern": r"Wireless Bar Code Reader\s*\(2D\)\s*\d+\s*\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            
            "C39": {
            "pattern": r"Laser Pointer[\s\S]*?\$?\s*([\d,]+\.\d{2})",
            "group": 1,
            "transform": lambda x: f"${x}"
            },

            "C41": {
            "pattern": r"Alphanumeric LED Bar[\s\S]*?\$?\s*([\d,]+\.\d{2})",
            "group": 1,
            "transform": lambda x: f"${x}"
            },

           "C43": {
            "pattern": r"RFID Badge Reader[\s\S]*?\$?\s*([\d,]+\.\d{2})",
            "group": 1,
            "transform": lambda x: f"${x}"
            },
           "C45": {
            "pattern": r"Spare Parts Kit[\s\S]*?\$?\s*([\d,]+\.\d{2})",
            "group": 1,
            "transform": lambda x: f"${x}"
            },
            "C47": {
                "pattern": r"Spare Parts Kit\s*[–-]\s*Level B[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C49": {
                "pattern": r"Spare Parts Kit\s*[–-]\s*Level C[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C51": {
                "pattern": r"Magnetic Badge Reader[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C53": {
                "pattern": r"Wireless Bar Code Reader\s*\(1D\)[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C55": {
                "pattern": r"Automatic Closing Door[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C57": {
                "pattern": r"Modula WMS Driver[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C59": {
                "pattern": r"Modula WMS Premium[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C61": {
                "pattern": r"Modula Link[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C63": {
                "pattern": r"Slotted Partitions L 857 H 120[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C65": {
                "pattern": r"Plain Dividers H 120 L 100[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C67": {
                "pattern": r"Plain Dividers H 120 L 140[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C69": {
                "pattern": r"Plain Dividers H 120 L 180[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C71": {
                "pattern": r"Plain Dividers H 120 L 220[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C73": {
                "pattern": r"Plain Dividers H 120 L 260[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            },
            "C75": {
                "pattern": r"Plain Dividers H 120 L 300[\s\S]*?\$?\s*([\d,]+\.\d{2})",
                "group": 1,
                "transform": lambda x: f"${x}"
            }
        }
    }            
    return mapping_rules



def insert_page_field(run):
    """
    Insert a Word PAGE field code into a run object

    Args:
        run: The run object to insert the field code into
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Create the field characters and instructions
    # This creates a proper Word field code that displays the current page number
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    # This is where a default value would go if Word fails to calculate
    # the page number, but we'll leave it empty since Word handles this well

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

def setup_document_headers_and_footers(doc, section, customer_name, proposal_num, current_date, page_num, skip_header=False, skip_footer=False):
    """
    Set up the headers and footers for a Word document section with a table-based layout

    Args:
        doc: The Word document
        section: The section to add headers/footers to
        customer_name: The customer name for the header
        proposal_num: The proposal number for the header
        current_date: The current date for the header
        page_num: The page number to display (used only for logging, not directly in document)
        skip_header: Whether to skip adding the header
        skip_footer: Whether to skip adding the footer
    """
    # Create unique headers for each section
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from datetime import datetime
    import logging

    logger = logging.getLogger(__name__)
    logger.debug(f"Setting header for page {page_num} with customer {customer_name}")

    # Ensure we have a valid date
    if not current_date:
        current_date = datetime.now().strftime('%B %d, %Y')

    if not skip_header:
        # Clear existing header content
        header = section.header
        for i in range(len(header.paragraphs)-1, -1, -1):
            p = header.paragraphs[i]
            p._element.getparent().remove(p._element)

        # Create a table for the header with 2 columns and specified width
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.0))
        header_table.allow_autofit = True

        # Set the width of the columns
        header_table.columns[0].width = Inches(4.0)  # Left column
        header_table.columns[1].width = Inches(2.9)  # Right column

        # Left cell content
        left_cell = header_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add customer info
        customer_run = left_para.add_run(f"{customer_name}\n")
        customer_run.font.name = 'Calibri'
        customer_run.font.size = Pt(11)

        title_run = left_para.add_run("Modula Vertical Lift Module – Proposal\n")
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(11)

        proposal_run = left_para.add_run(f"Proposal #{proposal_num}")
        proposal_run.font.name = 'Calibri'
        proposal_run.font.size = Pt(11)

        # Right cell content
        right_cell = header_table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add date
        date_run = right_para.add_run(f"{current_date}\n")
        date_run.font.name = 'Calibri'
        date_run.font.size = Pt(11)

        # Add "Page" text and field
        page_run = right_para.add_run("Page ")
        page_run.font.name = 'Calibri'
        page_run.font.size = Pt(11)

        # Insert dynamic page number
        insert_page_field(page_run)

        logger.debug(f"Added dynamic page field to header")

    if not skip_footer:
        # Clear any existing footer content but don't add page numbers
        footer = section.footer
        for i in range(len(footer.paragraphs)-1, -1, -1):
            p = footer.paragraphs[i]
            p._element.getparent().remove(p._element)

        # Add an empty paragraph to maintain consistent spacing
        footer.add_paragraph("")

def extract_customer_info_from_pdf(pdf_text):
    """
    Extract customer specific information from PDF

    Args:
        pdf_text: Full text content from PDF as a string

    Returns:
        Dictionary with customer info (name, location, date, proposal_number, contact_info)
    """
    import re

    # Normalize the text
    normalized_text = pdf_text.replace("\xa0", " ")
    normalized_text = re.sub(r"\s+", " ", normalized_text).strip()

    # Extract location
    location_patterns = [
        r"(?:Location|Address|Facility):\s*(.*?)(?:\n|$)",
        r"(?:in|at)\s+([A-Za-z\s,]+,\s+[A-Za-z]{2})",
        r"([A-Za-z\s,]+,\s+[A-Z]{2}\s+\d{5})"
    ]

    location_match = None
    for pattern in location_patterns:
        match = re.search(pattern, normalized_text, re.IGNORECASE)
        if match and match.group(1).strip():
            location_match = match.group(1).strip()
            customer_info["location"] = location_match
            break

    # Extract customer name as text immediately before location
    if location_match:
        escaped_location = re.escape(location_match)
        name_match = re.search(r"([A-Za-z\.\&\s]+?)\s+" + escaped_location, normalized_text)
        if name_match:
            customer_info["name"] = name_match.group(1).strip()
        else:
            customer_info["name"] = "Dart Aerospace"
    else:
        customer_info["name"] = "Dart Aerospace"


    # Extract location
    location_patterns = [
        r"(?:Location|Address|Facility):\s*(.*?)(?:\n|$)",
        r"(?:LOCATION|ADDRESS|FACILITY):\s*(.*?)(?:\n|$)",
        r"(?:in|at)\s+([A-Za-z\s,]+,\s+[A-Za-z]{2})",
        r"([A-Za-z\s,]+,\s+[A-Z]{2}\s+\d{5})"
    ]

    for pattern in location_patterns:
        match = re.search(pattern, pdf_text, re.IGNORECASE)
        if match and match.group(1).strip():
            customer_info["location"] = match.group(1).strip()
            break

    # Extract date
    date_patterns = [
        r"(?:Date|Dated):\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
        r"(?:DATE|DATED):\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
        r"(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
        r"(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}"
    ]

    for pattern in date_patterns:
        match = re.search(pattern, pdf_text, re.IGNORECASE)
        if match:
            customer_info["date"] = match.group(0).strip()
            break

    # Extract proposal number
    proposal_patterns = [
        r"(?:Proposal|Quote|Quotation|Offer)(?:\s+#|\s+No\.?|:|\s+Number:?)\s*([A-Z0-9\-\.]+)",
        r"(?:PROPOSAL|QUOTE|QUOTATION|OFFER)(?:\s+#|\s+NO\.?|:|\s+NUMBER:?)\s*([A-Z0-9\-\.]+)",
        r"Ref\.\s+Offer\s+No\.:\s*([A-Z0-9\-\.]+)",
        r"QLI-(\d+)"
    ]

    for pattern in proposal_patterns:
        match = re.search(pattern, pdf_text, re.IGNORECASE)
        if match:
            customer_info["proposal_number"] = match.group(1).strip()
            break

    # Extract model information
    model_patterns = [
        r"(?:Model|Unit|Machine)(?:\s+Type|\s+#|\s+No\.?|:|\s+Number:?)?\s*([A-Z0-9\-\.]+\s*\d+)",
        r"(?:MODEL|UNIT|MACHINE)(?:\s+TYPE|\s+#|\s+NO\.?|:|\s+NUMBER:?)?\s*([A-Z0-9\-\.]+\s*\d+)",
        r"Modula\s+([A-Z0-9\-\.]+\s*\d+)",
        r"ML25"
    ]

    for pattern in model_patterns:
        match = re.search(pattern, pdf_text, re.IGNORECASE)
        if match:
            customer_info["model"] = match.group(0).strip()
            break

    # Extract contact information
    contact_patterns = [
        r"(?:Contact|Attn|Attention|To the attention of):\s*(.*?)(?:\n|$)",
        r"(?:CONTACT|ATTN|ATTENTION|TO THE ATTENTION OF):\s*(.*?)(?:\n|$)",
        r"(?:Tel|Phone|Email):\s*(.*?)(?:\n|$)",
        r"(?:TEL|PHONE|EMAIL):\s*(.*?)(?:\n|$)"
    ]

    contact_info = []
    for pattern in contact_patterns:
        matches = re.findall(pattern, pdf_text, re.IGNORECASE)
        for match in matches:
            if match.strip():
                contact_info.append(match.strip())

    customer_info["contact_info"] = ", ".join(contact_info)

    # Extract unit height
    height_pattern = r"height\s+of\s+([\d\.]+)\s+inches\s+\(([\d\'\.\"]+)\)"
    height_match = re.search(height_pattern, pdf_text)
    if height_match:
        inches = height_match.group(1)
        feet = height_match.group(2)
        customer_info["unit_height"] = f"{inches} inches ({feet})"

    return customer_info


def process_pdf_to_word(pdf_path, word_path, excel_path= None, customer_info=None):
    """
    Process a PDF file and save extracted data to a Word document

    Args:
        pdf_path: Path to the PDF file
        word_path: Path to save the Word document
        excel_path: path to the Excel file to extract software info
        customer_info: Optional dictionary with customer information
    """
    import os
    import logging
    logger = logging.getLogger(__name__)

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create a new Word document
        doc = Document()

        # Set document-wide default font and line spacing
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.paragraph_format.line_spacing = 1.0

        # Set 1-inch margins for all sections
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Get access to sections for header/footer management
        sections = doc.sections

        try:
            # Extract text from the PDF - catching exceptions for each step
            text_content = extract_text_from_pdf(pdf_path)
        except Exception as text_e:
            logger.error(f"Error extracting text from PDF: {str(text_e)}")
            text_content = []
            # Add error information to the document
            doc.add_paragraph(f"Error extracting text: {str(text_e)}")

        try:
            # Extract tables from the PDF
            tables = extract_tables_from_pdf(pdf_path)
        except Exception as tables_e:
            logger.error(f"Error extracting tables from PDF: {str(tables_e)}")
            tables = []
            # Add error information to the document
            doc.add_paragraph(f"Error extracting tables: {str(tables_e)}")

        # Extract full text for parsing
        full_text = '\n'.join(text_content) if text_content else ""

        # Get customer info either from parameters or extract from PDF
        if not customer_info:
            customer_info = extract_customer_info_from_pdf(full_text)

        # Create mapping rules for extracting data
        mapping_rules = create_modula_mapping_rules()

        # Make sure we have a valid date
        from datetime import datetime
        current_date = customer_info.get('proposal_date', 
                        customer_info.get('date', datetime.now().strftime('%B %d, %Y')))
        if not current_date or current_date.lower() == 'none':
            current_date = datetime.now().strftime('%B %d, %Y')

        # Try to extract offer number
        offer_number = customer_info.get("proposal_number", "")
        offer_rule = mapping_rules["Primary Option"]["B8"]
        if "pattern" in offer_rule:
            import re
            offer_match = re.search(offer_rule["pattern"], full_text)
            if offer_match:
                offer_number = offer_match.group(1)

        # Try to extract number of trays
        num_trays = "22"  # Default from provided text
        trays_rule = mapping_rules["Primary Option"]["B9"]
        if "pattern" in trays_rule:
            import re
            trays_match = re.search(trays_rule["pattern"], full_text)
            if trays_match:
                num_trays = trays_match.group(1)

        # Extract tray dimensions and other specifications
        tray_dimensions = format_tray_dimensions(full_text)
        
        #new brian code 
        # Extract Storage Area
        storage_area = "N/A"  # default
        storage_area_rule = mapping_rules["Primary Option"]["B13"]  # Storage Area
        if "pattern" in storage_area_rule:
            import re
            area_match = re.search(storage_area_rule["pattern"], full_text, re.IGNORECASE)
            if area_match:
                storage_area = area_match.group(storage_area_rule["group"])
                    
        # Extract Storage Volume
        storage_volume = "N/A"  # default
        storage_volume_rule = mapping_rules["Primary Option"]["B12"]  # Storage Volume
        if "pattern" in storage_volume_rule:
            import re
            volume_match = re.search(storage_volume_rule["pattern"], full_text, re.IGNORECASE)
            if volume_match:
                storage_volume = volume_match.group(storage_volume_rule["group"])

        # Extract Unit Footprint
        import re

        unit_footprint = "N/A"  # default
        unit_footprint_rule = mapping_rules["Primary Option"]["B14"]  # Unit Footprint

        # Ensure both pattern and group exist before trying to extract
        if unit_footprint_rule and "pattern" in unit_footprint_rule and "group" in unit_footprint_rule:
            pattern = unit_footprint_rule["pattern"]
            group_num = unit_footprint_rule["group"]

            # Clean full_text to avoid hidden line breaks or spacing issues
            cleaned_text = re.sub(r'\s+', ' ', full_text.strip())

            # Perform regex search
            footprint_match = re.search(pattern, cleaned_text, re.IGNORECASE)
            if footprint_match:
                try:
                    unit_footprint = footprint_match.group(group_num).strip()
                except IndexError:
                    unit_footprint = "N/A"  # fallback if group not found

        # Extract Height
        height = "N/A"  # default
        height_rule = mapping_rules["Primary Option"]["B11"]  # Height
        if "pattern" in height_rule:
            import re
            height_match = re.search(height_rule["pattern"], full_text, re.IGNORECASE)
            if height_match:
                height = height_match.group(height_rule["group"])


        # Try to extract tray dimensions (Width and Depth)
        # Default tray dimensions
        tray_dimensions = "Width: N/A x Depth: N/A"  

        tray_rule = mapping_rules["Primary Option"]["B10"]

        # Check if a custom handler exists
        if "custom_handler" in tray_rule:
            tray_dimensions = tray_rule["custom_handler"](full_text)  # call format_tray_dimensions function
        #     tray_dimensions = f"Width: {width}'' x Depth: {depth}''"



        # ======== COVER PAGE ========
        # First section - for cover page
        current_section = doc.sections[0]  # First section (0-indexed)
        
        
        name_to_use = customer_info.get('customer_name', customer_info.get('name', 'Dart Aerospace'))
        if name_to_use.strip().lower() == "customer":
            name_to_use = "Dart Aerospace"

        setup_document_headers_and_footers(
            doc,
            current_section,
            name_to_use,
            customer_info.get('proposal_number', 'N/A'),
            current_date,
            1
        )

        # Set up header with page number 1
        # setup_document_headers_and_footers(doc, current_section, 
        #                                    customer_info.get('customer_name', customer_info.get('name', 'Dart Aerospace')),
        #                                    customer_info.get('proposal_number', 'N/A'),
        #                                    current_date, 1)

        # Add PIH Logo
        logo_paragraph = doc.add_paragraph("")
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Path to PIH logo
        pih_logo_path = os.path.join('static', 'images', 'pihlogo.png')

        if os.path.exists(pih_logo_path):
            # Add the actual PIH logo
            try:
                logo_paragraph.add_run().add_picture(pih_logo_path, width=Inches(3.5))
                logger.debug(f"Added PIH logo from {pih_logo_path}")
            except Exception as logo_e:
                logger.error(f"Error adding PIH logo: {str(logo_e)}")
                # Fallback to text if there's an error
                logo_run = logo_paragraph.add_run("[PIH LOGO PLACEHOLDER]")
                logo_run.font.size = Pt(14)
                logo_run.bold = True
        else:
            # Fallback if logo file is not found
            logo_run = logo_paragraph.add_run("[PIH LOGO PLACEHOLDER]")
            logo_run.font.size = Pt(14)
            logo_run.bold = True

        # Add PIH address
        address_paragraph = doc.add_paragraph()
        address_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_run = address_paragraph.add_run("Pacific Integrated Handling\n10215 Portland Ave E\nTacoma, WA 98455")
        address_run.font.size = Pt(11)

        # Add "Presents to:" text
        presents_paragraph = doc.add_paragraph()
        presents_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        presents_run = presents_paragraph.add_run("Presents to:")
        presents_run.italic = True
        presents_run.font.size = Pt(11)

        # Add customer logo (prefer uploaded logo if provided, otherwise fallback to PDF crop or placeholder)
        customer_logo_paragraph = doc.add_paragraph()
        customer_logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Prefer explicitly uploaded customer logo if available
        uploaded_logo_path = (customer_info or {}).get('logo_path')
        if uploaded_logo_path and os.path.exists(uploaded_logo_path):
            try:
                run = customer_logo_paragraph.add_run()
                run.add_picture(uploaded_logo_path, width=Inches(2.5))
                logger.info(f"Added uploaded customer logo from {uploaded_logo_path}")
            except Exception as logo_e:
                logger.error(f"Error adding uploaded customer logo: {str(logo_e)}")
        else:
            # Fallback: attempt to crop a logo-like region from the first page of the source PDF
            try:
                import fitz  # PyMuPDF
                import tempfile, os
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                # Open the PDF and take the first page (index 0)
                pdf_doc = fitz.open(pdf_path)
                page = pdf_doc[0]
                rect = page.rect  

                # Define crop area for top-left corner
                crop_width = rect.width * 0.25
                crop_height = rect.height * 0.12  

                # Margins: adjust top & right only
                margin_top = 15    # crop a bit more from the top
                margin_right = 15  # crop a bit more from the right

                crop_rect = fitz.Rect(
                    rect.x0,                          # keep left as is
                    rect.y0 + margin_top,             # push down from the top
                    rect.x0 + crop_width - margin_right,  # trim more from the right
                    rect.y0 + crop_height             # keep bottom as is
                )

                pix = page.get_pixmap(dpi=300, clip=crop_rect)
        
                # Render cropped portion
                pix = page.get_pixmap(dpi=300, clip=crop_rect)

                # Save cropped image temporarily
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                    temp_path = temp_img.name
                pix.save(temp_path)
                pdf_doc.close()

                # Insert into Word document (same logic as your logo code)
                if os.path.exists(temp_path):
                    customer_logo_paragraph = doc.add_paragraph()
                    customer_logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = customer_logo_paragraph.add_run()
                    run.add_picture(temp_path, width=Inches(2.5))  # scale logo size similar to original
                    os.unlink(temp_path)
                    logger.info("Extracted first page cropped image as customer logo.")
                else:
                    # Fallback to placeholder text if something fails
                    cust_name = customer_info.get('customer_name', customer_info.get('name', 'Customer'))
                    customer_logo_paragraph = doc.add_paragraph()
                    customer_logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    customer_logo_run = customer_logo_paragraph.add_run(f"[{cust_name} LOGO]")
                    customer_logo_run.font.size = Pt(24)
                    customer_logo_run.bold = True

            except Exception as e:
                logger.error(f"Error extracting logo from PDF: {str(e)}")
                cust_name = customer_info.get('customer_name', customer_info.get('name', 'Customer'))
                customer_logo_paragraph = doc.add_paragraph()
                customer_logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                customer_logo_run = customer_logo_paragraph.add_run(f"[{cust_name} LOGO]")
                customer_logo_run.font.size = Pt(24)
                customer_logo_run.bold = True


        # Add location and quote
        location_paragraph = doc.add_paragraph()
        location_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add fallback for location
        location = customer_info.get('location', 'Location Not Specified')
        location_run = location_paragraph.add_run(f"{location}")
        location_run.font.size = Pt(11)

        quote_paragraph = doc.add_paragraph()
        quote_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        quote_run = quote_paragraph.add_run("Quote")
        quote_run.font.size = Pt(14)
        quote_run.bold = True

        # Add model description
        model_desc_paragraph = doc.add_paragraph()
        model_desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Get model from 'model' or fall back to 'vlm_model' if available
        model_desc = customer_info.get('model', customer_info.get('vlm_model', 'Modula VLM'))
        model_desc_run = model_desc_paragraph.add_run(f"{model_desc}")
        model_desc_run.font.size = Pt(12)

        # Add VLM image for cover page: prefer user-selected image, fallback to default asset
        vlm_image_path = (customer_info or {}).get('vlm_image_path') or os.path.join('static', 'images', 'vlm machine.png')

        if vlm_image_path and os.path.exists(vlm_image_path):
            # Add the actual PIH logo
            try:
                vlm_image_paragraph = doc.add_paragraph()
                vlm_image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Use a larger width for the VLM image (cover page feature)
                vlm_image_paragraph.add_run().add_picture(vlm_image_path, width=Inches(1.5))
                logger.debug(f"Added vlm image from {vlm_image_path}")
            except Exception as logo_e:
                logger.error(f"Error adding vlm image: {str(logo_e)}")
                # Fallback to text if there's an error
                logo_run = vlm_image_paragraph.add_run("[VLM IMAGE PLACEHOLDER]")
                logo_run.font.size = Pt(14)
                logo_run.bold = True
        else:
            # Fallback if logo file is not found
            logo_run = vlm_image_paragraph.add_run("[VLM IMAGE  PLACEHOLDER]")
            logo_run.font.size = Pt(14)
            logo_run.bold = True

        # Add contact info at bottom
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Default contact values
        contact_name = customer_info.get('contact_name', 'Josh Jancola')
        contact_phone = customer_info.get('contact_phone', '253.500.4193')
        contact_office = customer_info.get('contact_office', '888.550.5888')
        contact_email = customer_info.get('contact_email', 'joshjancola@pacificintegrated.com')

        contact_run = contact_paragraph.add_run(
            f"{contact_name}\n"
            f"Cell: {contact_phone}\n"
            f"Office: {contact_office}\n"
            f"{contact_email}"
        )
        contact_run.font.size = Pt(10)

        # Add page break after cover page
        doc.add_page_break()

        # ======== EXECUTIVE SUMMARY PAGE ========
        # ======== PAGE 2: EXECUTIVE SUMMARY ========
        # Create a new section for the executive summary page
        doc.add_section()

        # Get customer and date info
        # Get customer and date info (robust version)
        client_name = name_to_use
        proposal_num = (customer_info.get('proposal_number') or customer_info.get('proposal_num') or '').strip() 
        current_date = (customer_info.get('proposal_date') or customer_info.get('date') or '').strip() 


        # Set up header and footer for page 2
        current_section = doc.sections[1]  # Second section (0-indexed)
        setup_document_headers_and_footers(doc, current_section, client_name, proposal_num, current_date, 2)

        # Define client variables for the summary - use customer_name for backward compatibility
        facility_location = customer_info.get('location', 'Hillsboro, OR')
        model_value = customer_info.get('model', customer_info.get('vlm_model', 'MODULA®LIFT NG VLM, Model ML25'))
        num_trays = customer_info.get('tray_quantity', customer_info.get('num_trays', '22'))

        # Add recipient name (optional)
        recipient_name = customer_info.get('contact_person', 'Jonathon Johnson')
        recipient_para = doc.add_paragraph()
        recipient_para.add_run(f"{recipient_name},").bold = False
        
        # new brian code
        # if excel_path:
        #     software_name = extract_software_name_from_excel(excel_path)
        # else:
        #     software_name = "Modula WMS Premium"
              # fallback
        
        # new brian code 
        # Extract software name
        # Extract software name from Excel
        if customer_info and customer_info.get('excel_path'):
            excel_path = customer_info.get('excel_path')
            # Normalize the path to handle both forward and backward slashes
            excel_path = os.path.normpath(excel_path)
        extract_optional_items_from_excel(excel_path)


        # Add Executive Summary - main content in properly formatted paragraphs
        # First paragraph - cleanroom storage solution
        para1 = doc.add_paragraph()
        para1_text = f"Pacific Integrated Handling thanks you for the opportunity to provide {client_name} with the following automated storage solution for the cleanroom at the {facility_location}. The following proposal provides an in-depth description of the {model_value} we are proposing."
        para1.add_run(para1_text).bold = False

        # Paragraph 2 - Modula VLM benefits
        para2 = doc.add_paragraph()
        para2_text_part1 = "This "
        para2_text_part2 = f" Vertical Lift Module (VLM) has been specified to meet {client_name}'s storage requirements and will aid in increasing storage density and capacity. Computerized management also enables all the logistics processes to be monitored and optimized. This vertical system will free up valuable floor space, minimize search time, and increase productivity utilizing the latest technology available."
        para2.add_run(para2_text_part1).bold = False
        para2.add_run("Modula").bold = True
        para2.add_run(para2_text_part2).bold = False

        # Paragraph 3 - Storage capacity details
        para3 = doc.add_paragraph()
        para3_text = f"The VLM will have a total storage capacity of {storage_volume} sq ft. and {storage_area} cubic ft of volume in a very small footprint of {unit_footprint} sq ft. While on-site, I completed a survey of the existing materials in the cleanroom and concluded that (with a 50% buffer) only 136.80 cubic ft is needed for storage. This will give {client_name} plenty of room for growth or flexibility in what is kept within the VLM. A safety photo eye curtain system will protect operators during mechanical movements. All safety devices and machine designs meet CE standards. In addition, tool accessible doors on the side of the unit will allow service personnel access to the units for scheduled maintenance. The VLM will come with 2 years of parts and labor warranty as well as a PIH exclusive 3 scheduled maintenance visits."
        para3.add_run(para3_text).bold = False

        # Paragraph 4 - Tray information
        para4 = doc.add_paragraph()
        para4_text = f"Additionally, the VLM have a height of {height} inches and will contain {num_trays} slotted trays - all of which are sized as {tray_dimensions} and will be capable of holding up to 551 lbs. Slotted trays will allow {client_name} to use Modula dividers to create storage cells within the trays to organize smaller parts."
        para4.add_run(para4_text).bold = False
        
        # Paragraph 5 - Software details
        print(f"Software = {VALUE}"*1)
        para5 = doc.add_paragraph()
        para5_text = f"This proposal includes software installation of {VALUE}, which is comprehensive software and will allow for the incorporation of bar code scanning, ERP integration, and future implementation of put-to-light systems, per the discussion on site."
        para5.add_run(para5_text).bold = False

        # Paragraph 6 - ISO7 readiness
        para6 = doc.add_paragraph()
        para6_text = f"The Modula VLM is considered ISO7-Ready and will be the ideal fit for the cleanroom environment at {client_name}."
        para6.add_run(para6_text).bold = False

        # Paragraph 7 - PIH experience and conclusion
        para7 = doc.add_paragraph()
        para7_text = f"PIH has decades of experience with automated material handling systems. Our factory trained and certified technicians have a meticulous attention to detail and a reputation for the highest level of customer service and satisfaction, not only during the installation, but also on the following scheduled maintenance visits. Without hesitation I can say that PIH is the best partner for {client_name} when it comes to your present and future automated high-density storage and retrieval system needs."
        para7.add_run(para7_text).bold = False

        # Add contact information
        contact_para = doc.add_paragraph()
        contact_name = customer_info.get('contact_name', 'Josh Jancola')
        contact_email = customer_info.get('contact_email', 'joshjancola@pacificintegrated.com')
        contact_phone = customer_info.get('contact_phone', '253.500.4193')

        contact_para.add_run(f"{contact_name}\n").bold = False
        contact_para.add_run("Pacific Integrated Handling\n").bold = False
        contact_para.add_run(f"Cell: {contact_phone}\n").bold = False
        contact_para.add_run(f"{contact_email}").bold = False

        # We're eliminating the extra content from page 2 to ensure executive summary fits

        # Add a page break before  (now page 3, since we removed Project Overview)
        doc.add_page_break()

        # ========  (PAGE 3) ========
        # Create a new section for page 3 (was previously page 4)
        doc.add_section()

        # Get customer and date info
        # Get customer and date info
        # Get customer and date info (robust version)
        client_name = (customer_info.get('customer_name') or customer_info.get('name') or '').strip() 
        proposal_num = (customer_info.get('proposal_number') or customer_info.get('proposal_num') or '').strip() 
        current_date = (customer_info.get('proposal_date') or customer_info.get('date') or '').strip() 



        # Set up header and footer for page 3
        current_section = doc.sections[2]  # Third section (0-indexed)
        setup_document_headers_and_footers(doc, current_section, client_name, proposal_num, current_date, 3)


        # Add "Why Pacific Integrated Handling?" section
        why_pih_title = doc.add_heading('Why Pacific Integrated Handling?', level=1)
        why_pih_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add the content with proper bullets and sub-bullets

        # Using properly formatted bullet points with nested levels

        # Benefit 1 - Tier 1 dealer
        p1 = doc.add_paragraph()
        p1.style = 'List Bullet'
        p1.add_run("PIH is the Tier 1 dealer for Modula in the Pacific Northwest, which comes with the following exclusive benefits:")

        # Sub-bullets for benefit 1
        p1_1 = doc.add_paragraph()
        p1_1.style = 'List Bullet 2'  # Second level bullet (circle)
        p1_1.add_run("PIH boasts our own team of factory-trained service and support technicians.")

        p1_2 = doc.add_paragraph()
        p1_2.style = 'List Bullet 2'
        p1_2.add_run("Certified Technicians are all outfitted with fully stocked service vans.")

        p1_3 = doc.add_paragraph()
        p1_3.style = 'List Bullet 2'
        p1_3.add_run("Direct partnership and engagement with the Modula support team.")

        # Benefit 2 - Scheduled maintenance
        p2 = doc.add_paragraph()
        p2.style = 'List Bullet'
        p2.add_run("PIH offers 3 scheduled maintenance visits throughout the 2-year warranty period.")

        # Sub-bullet for benefit 2
        p2_1 = doc.add_paragraph()
        p2_1.style = 'List Bullet 2'
        p2_1.add_run("Modula and their partners only require and perform 1 visit.")

        # Benefit 3 - Largest dealer
        p3 = doc.add_paragraph()
        p3.style = 'List Bullet'
        p3.add_run("PIH is the largest Modula VLM dealer in the PNW, actively servicing dozens of clients across Oregon and the rest of the Western seaboard.")

        # Sub-bullet for benefit 3
        p3_1 = doc.add_paragraph()
        p3_1.style = 'List Bullet 2'
        p3_1.add_run("This high volume of VLMs under PIH service agreements directly results in highly trained and experienced technicians.")

        # Remaining benefits
        p4 = doc.add_paragraph()
        p4.style = 'List Bullet'
        p4.add_run("PIH offers a full turnkey solution without the need for outsourcing labor.")

        p5 = doc.add_paragraph()
        p5.style = 'List Bullet'
        p5.add_run("PIH stocks over $250,000 worth of spare parts in case of mechanical emergencies with our valued customers.")

        p6 = doc.add_paragraph()
        p6.style = 'List Bullet'
        p6.add_run("PIH's Systems Division is singularly focused on high density automated storage and retrieval systems and has been since 1981.")

        p7 = doc.add_paragraph()
        p7.style = 'List Bullet'
        p7.add_run("PIH has a fully staffed CAD design team dedicated to supporting the conceptualization and installation of VLMs in our customers' facilities.")
        
        
        # Add "Why Modula?" section
        why_modula_title = doc.add_heading('Why Modula?', level=1)
        why_modula_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        modula_benefits_text = """- Components of the Modula VLM are off the shelf and not proprietary (greater
  availability at lower cost).
- One-time cost for software. Includes software version updates at no additional cost.
- VLM control panel (co-pilot) is at an ergonomic height and in a location that minimizes
  risk of safety barrier tripping.
- Modula has the most advanced anti-seismic system on the market today.
- Modula WMS has a highly intuitive, modern user interface.
- Modula VLMs are the only VLMs which can meet "Made in USA" and "Buy America"
  requirements.
- VLM software can interface with Corridor or any other ERP/tool management system to ensure
  accountability of items both inside and outside the machine. It also can operate as a robust,
  standalone system."""
        doc.add_paragraph(modula_benefits_text)

        # Add a page break before Why PIH section
        doc.add_page_break()

       # ======== PAGE 4: SCREENSHOT FROM PDF PAGE 12 ========
        doc.add_section()
        section_index = len(doc.sections) - 1
        current_section = doc.sections[section_index]
        current_page_num = 4
        setup_document_headers_and_footers(doc, current_section, client_name, proposal_num, current_date, current_page_num)
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[11]  # PDF page 5
            rect = page.rect
            crop_top_px=100
            crop_bottom_px=100
            crop_rect = fitz.Rect(rect.x0, rect.y0 + crop_top_px, rect.x1, rect.y1 - crop_bottom_px)
            pix = page.get_pixmap(dpi=1000, clip=crop_rect)
              
            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name
            pix.save(temp_path)
            pdf_doc.close()

            if os.path.exists(temp_path):
                section = doc.sections[-1]

                available_width = section.page_width - section.left_margin - section.right_margin

                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()

                # Only set width to fit page, height adjusts automatically
                run.add_picture(temp_path, width=available_width)

                # Remove extra spacing
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

                # Reduce section margins if needed
                section.top_margin = Inches(1.9)
                section.bottom_margin = Inches(0.5)


                logger.info("Added cropped PDF page 5 image to Word document below header and centered")
                os.unlink(temp_path)
            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF for machine specifications.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")

        doc.add_page_break()



        # ======== PAGE 5: SCREENSHOT FROM PDF PAGE 13 ========
        doc.add_section()
        section_index = len(doc.sections) - 1
        current_section = doc.sections[section_index]
        current_page_num = 5
        setup_document_headers_and_footers(doc, current_section, client_name, proposal_num, current_date, current_page_num)
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[12]  # PDF page 5
            rect = page.rect
            crop_top_px=100
            crop_bottom_px=30
            crop_rect = fitz.Rect(rect.x0, rect.y0 + crop_top_px, rect.x1, rect.y1 - crop_bottom_px)
            pix = page.get_pixmap(dpi=800, clip=crop_rect)
              
            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name
            pix.save(temp_path)
            pdf_doc.close()

            if os.path.exists(temp_path):
                section = doc.sections[-1]
                # Compute max available width
                available_width = section.page_width - section.left_margin - section.right_margin

                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()

                # Only set width to fit page, height adjusts automatically
                run.add_picture(temp_path, width=available_width)

                # Remove extra spacing
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

                # Reduce section margins if needed
                section.top_margin = Inches(1.9)
                section.bottom_margin = Inches(0.5)

                logger.info("Added cropped PDF page 5 image to Word document below header and centered")
                os.unlink(temp_path)
            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF for machine specifications.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")

        doc.add_page_break()
 
                # ======== PAGE 6: SCREENSHOT FROM PDF PAGE 14 ========
        doc.add_section()
        section_index = len(doc.sections) - 1
        current_section = doc.sections[section_index]
        current_page_num = 6
        setup_document_headers_and_footers(doc, current_section, client_name, proposal_num, current_date, current_page_num)
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[13]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 80)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()

            if os.path.exists(temp_path):
                section = doc.sections[-1]

                # keep margins respected
                available_width = section.page_width - section.left_margin - section.right_margin

                # adjust section top margin so text starts below header
                section.top_margin = Inches(2.0)  # increase if still too close to header

                # create a centered paragraph for the image
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER   # center the image container

                run = para.add_run()
                run.add_picture(temp_path, width=available_width)

                # formatting inside paragraph (extra gap above image)
                para_format = para.paragraph_format
                para_format.space_before = Pt(12)  # additional push down from header

                logger.info("Added cropped PDF page 5 image to Word document below header and centered")
                os.unlink(temp_path)
            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF for machine specifications.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")

        doc.add_page_break()
        
        
              # ======== PAGE 7: SCREENSHOT FROM PDF PAGE 15 ========
        doc.add_section()
        section_index = len(doc.sections) - 1
        current_section = doc.sections[section_index]
        current_page_num = 7
        setup_document_headers_and_footers(doc, current_section, client_name, proposal_num, current_date, current_page_num)
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[14]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()

            if os.path.exists(temp_path):
                section = doc.sections[-1]

                # keep margins respected
                available_width = section.page_width - section.left_margin - section.right_margin

                # adjust section top margin so text starts below header
                section.top_margin = Inches(2.0)  # increase if still too close to header

                # create a centered paragraph for the image
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER   # center the image container

                run = para.add_run()
                run.add_picture(temp_path, width=available_width)

                # formatting inside paragraph (extra gap above image)
                para_format = para.paragraph_format
                para_format.space_before = Pt(12)  # additional push down from header

                logger.info("Added cropped PDF page 5 image to Word document below header and centered")
                os.unlink(temp_path)
            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF for machine specifications.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")

        doc.add_page_break()
          
        # ======== FINAL PAGE: PRICING TABLE ========
        # Create a new section for the pricing table page
        doc.add_section()
        section_index = len(doc.sections) - 1
        current_section = doc.sections[section_index]
        current_page_num = 7
        setup_document_headers_and_footers(doc, current_section, client_name, proposal_num, current_date, current_page_num)

        # Add pricing table section header
        system_investment_title = doc.add_heading(f'System Investment: 1 {model_value}', level=1)
        system_investment_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add "Unit Pricing" subheader
        unit_pricing_para = doc.add_paragraph()
        unit_pricing_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        unit_pricing_run = unit_pricing_para.add_run("Unit Pricing")
        unit_pricing_run.font.size = Pt(14)
        unit_pricing_run.bold = True
        unit_pricing_run.font.color.rgb = RGBColor(65, 105, 225)  # Royal blue

        # Create pricing table with 4 columns: Items, Price Each, Qty, and Total Price
        pricing_table = doc.add_table(rows=1, cols=4)
        pricing_table.style = 'Table Grid'

        # Set up table formatting based on ChatGPT's suggestion
        pricing_table.autofit = False  # Disable autofit to allow manual width control

        # Set absolute column widths based on a standard 6.5 inch content width
        # (standard letter paper with 1-inch margins on each side)
        width_items = Inches(4.0)   # ~65% of available width for items column
        width_price = Inches(0.85)  # ~12% of available width
        width_qty = Inches(0.65)    # ~10% of available width 
        width_total = Inches(1.0)   # ~13% of available width

        # Set the column widths
        pricing_table.columns[0].width = width_items
        pricing_table.columns[1].width = width_price
        pricing_table.columns[2].width = width_qty
        pricing_table.columns[3].width = width_total

        # Apply explicit width to each cell in the header row
        for i, cell in enumerate(pricing_table.rows[0].cells):
            if i == 0:
                cell.width = width_items
            elif i == 1:
                cell.width = width_price
            elif i == 2:
                cell.width = width_qty
            elif i == 3:
                cell.width = width_total

        # Style the header row
        header_cells = pricing_table.rows[0].cells
        header_cells[0].text = "Items"
        header_cells[1].text = "Price Ea."
        header_cells[2].text = "Qty."
        header_cells[3].text = "Prices"

        # Center the column headers and make them bold
        for i, cell in enumerate(header_cells):
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell_para.runs:
                run.bold = True

        # Get line items from the database if provided (from the review page selections)
        line_items = []
        if hasattr(customer_info, 'get') and customer_info.get('line_items'):
            line_items = customer_info.get('line_items')

        # If no line items provided, create default ones based on extracted info
        if not line_items:
            # Extract basic pricing information from customer_info
            base_price = customer_info.get('base_price', 100000.00)
            software_price = customer_info.get('software_price', 0.00)
            installation_price = customer_info.get('installation_price', 0.00)
            seismic_price = customer_info.get('seismic_price', 5500.00)
            freight_price = customer_info.get('freight_price', 16901.00)

            # Calculate tray quantities, ensuring integer division
            try:
                num_trays_int = int(num_trays)
                half_trays = num_trays_int // 2  # Integer division
                # If we have an odd number of trays, add the extra one to the first type
                first_half = half_trays + (num_trays_int % 2)
                second_half = half_trays
            except (ValueError, TypeError):
                # Default to 26 total trays (13 of each) if num_trays can't be parsed
                first_half = 13
                second_half = 13

            # Create default line items structure
            line_items = [
                # VLM Configuration section
                {"category": "VLM", "description": f"{model_value} With Single-Tray Presentation", 
                 "price_each": base_price, "quantity": 1, "is_included": False, "is_section_header": False},
                {"category": "VLM", "description": "Single Bay & Internal WorkStation", 
                 "price_each": 0, "quantity": 1, "is_included": True, "is_section_header": False},
                {"category": "VLM", "description": f"Tray size: 161.41\" W x 25.75\" D Tray Height 2.76\"", 
                 "price_each": 0, "quantity": first_half, "is_included": True, "is_section_header": False},
                {"category": "VLM", "description": f"Tray size: 161.41\" W x 25.75\" D Tray Height 4.72\"", 
                 "price_each": 0, "quantity": second_half, "is_included": True, "is_section_header": False},
                {"category": "VLM", "description": "Co-Pilot Touchscreen Operator Console", 
                 "price_each": 0, "quantity": 1, "is_included": True, "is_section_header": False},
                {"category": "VLM", "description": "Telephone Support – 8 Hours/Day, 5 Days/Week – 2 Years", 
                 "price_each": 0, "quantity": 1, "is_included": True, "is_section_header": False},

                # Software section
                {"category": "Software", "description": "Software", 
                 "price_each": 0, "quantity": 0, "is_included": False, "is_section_header": True},
                {"category": "Software", "description": "Modula WMS Base Software (includes on-site installation and training)", 
                 "price_each": 0, "quantity": 1, "is_included": True, "is_section_header": False},
        {"category": "Software", "description": "PIH On-Site Software Support", 
         "price_each": "FILL IN", "quantity": 1, "is_included": False, "is_section_header": False},

                # Installation section
                {"category": "Installation", "description": "Installation and Warranty", 
                 "price_each": 0, "quantity": 0, "is_included": False, "is_section_header": True},
                {"category": "Installation", "description": "Mechanical Installation (Includes Equipment Rentals)", 
                 "price_each": 0, "quantity": 1, "is_included": True, "is_section_header": False},
                {"category": "Installation", "description": "PIH Project Management", 
                 "price_each": 0, "quantity": 1, "is_included": True, "is_section_header": False},
                {"category": "Installation", "description": "Parts and Labor Warranty (2 Years)", 
                 "price_each": 0, "quantity": 1, "is_included": True, "is_section_header": False},
                {"category": "Installation", "description": "Seismic and Anchoring Calculations and Certification (Required)", 
                 "price_each": seismic_price, "quantity": 1, "is_included": False, "is_section_header": False},
                {"category": "Installation", "description": "PIH Exclusive: 3 Scheduled Maintenance Visits (Modula only requires 1)", 
                 "price_each": 0, "quantity": 1, "is_included": True, "is_section_header": False},

                # Freight section
                {"category": "Freight", "description": "Freight", 
                 "price_each": 0, "quantity": 0, "is_included": False, "is_section_header": True},
                {"category": "Freight", "description": f"Freight to {facility_location} (2 Calistoga Trucks, Side Load)", 
                 "price_each": freight_price, "quantity": 1, "is_included": False, "is_section_header": False},
            ]

        # Create a function to format currency values
        def format_currency(value):
            if value is None:
                return ""
            if value == 0:
                return "Included"
            return f"${value:,.2f}"

        # Function to format the display of quantity
        def format_quantity(qty):
            if qty is None:
                return ""
            if isinstance(qty, (int, float)):
                if qty == int(qty):  # If it's a whole number
                    return str(int(qty))
                return str(qty)  # Otherwise return with decimal
            return str(qty)  # For anything else, convert to string

        # Track current category for section headers
        current_category = None
        total_price = 0

        # Add line items to the table
        for item in line_items:
            # Skip optional items and section headers marked as optional,
            # they'll go in a separate table or be excluded
            if item.get('is_optional', False):
                continue

            # Items that should never be marked as included
            never_included_items = ["Virtual Partitioning Logic", "Operator Login Functionality"]
            item_description = item.get('description', '')

            # Force these items to not be included
            if any(never_item.lower() in item_description.lower() for never_item in never_included_items):
                item['is_included'] = False

            # Add a spacer row before each new section (category)
            if current_category != item.get('category'):
                current_category = item.get('category')

            # Add the item row
            row = pricing_table.add_row()
            cells = row.cells

            # Apply column widths to the new row's cells to maintain proportions
            cells[0].width = width_items
            cells[1].width = width_price
            cells[2].width = width_qty
            cells[3].width = width_total

            # Handle section headers differently
            is_section_header = item.get('is_section_header', False)
            if is_section_header:
                cells[0].text = item.get('description', '')
                cells[0].merge(cells[3])  # Merge all cells for section header
                header_para = cells[0].paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in header_para.runs:
                    run.bold = True
                continue

            # Regular item row
            cells[0].text = item.get('description', '')

            # Price each cell
            price_each = item.get('price_each', 0)
            is_included = item.get('is_included', False)

            if is_included:
                cells[1].text = "Included"
            else:
                cells[1].text = format_currency(price_each)

            # Quantity cell
            quantity = item.get('quantity', 1)
            cells[2].text = format_quantity(quantity) + (" lot" if quantity == 1 and item.get('category') in ['Installation', 'Freight'] else "")

            # Price total cell
            if is_included:
                cells[3].text = "Included"
            else:
                price_total = price_each * quantity
                total_price += price_total
                cells[3].text = format_currency(price_total)

            # Align price and quantity columns to center
            for i in range(1, 4):
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add total row
        total_row = pricing_table.add_row()
        total_cells = total_row.cells

        # Apply column widths to the total row to maintain proportions
        total_cells[0].width = width_items
        total_cells[1].width = width_price
        total_cells[2].width = width_qty
        total_cells[3].width = width_total

        # Merge the first 3 cells for "Total" text
        total_cells[0].merge(total_cells[2])
        total_para = total_cells[0].paragraphs[0]
        total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_run = total_para.add_run("Total:")
        total_run.bold = True
        total_run.font.size = Pt(12)

        # Add the total price in the last cell
        total_cells[3].text = format_currency(total_price)
        total_price_para = total_cells[3].paragraphs[0]
        total_price_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in total_price_para.runs:
            run.bold = True
            run.font.size = Pt(12)


        #NEW CODE 
        from pathlib import Path
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # ======== OPTIONAL ITEMS SECTION ========
        doc.add_section()
        current_section = doc.sections[-1]

        # Reduce top margin
        current_section.top_margin = Inches(0.5)
        current_section.header_distance = Inches(0.25)

        # Setup headers/footers
        current_page_num = 9
        setup_document_headers_and_footers(doc, current_section, client_name, proposal_num, current_date, current_page_num)

        # Create heading
        heading = doc.add_heading('Optional Items & Accessories Pricing', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)

        for run in heading.runs:
            run.font.size = Pt(14)

        # Add explanatory text
        explanation_para = doc.add_paragraph()
        explanation_para.add_run(
            "All items shown in the following pricing table are optional and are available for an additional fee over base solution pricing shown on the previous pages."
        )

        # Extract optional items from uploaded Excel file if available
        optional_items_data = []
        
        if customer_info and customer_info.get('excel_path'):
            excel_path = customer_info.get('excel_path')
            # Normalize the path to handle both forward and backward slashes
            excel_path = os.path.normpath(excel_path)
            logger.info(f"Excel path found: {excel_path}")
            try:
                if os.path.exists(excel_path):
                    logger.info(f"Excel file exists, extracting data...")
                    optional_items_data = extract_optional_items_from_excel(excel_path)
                    logger.info(f"Extracted {len(optional_items_data)} optional items from Excel file")
                else:
                    logger.warning(f"Excel file not found at {excel_path}")
            except Exception as e:
                logger.error(f"Error extracting optional items from Excel: {str(e)}")
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")
                optional_items_data = []
        else:
            logger.warning("No Excel path found in customer_info")

        # Create optional items table if we have data
        if optional_items_data:
            # Create table with proper columns
            options_table = doc.add_table(rows=1, cols=3)
            options_table.style = 'Table Grid'
            options_table.autofit = False

            # Set column widths
            width_items = Inches(4.0)
            width_price_est = Inches(1.5)
            width_qty = Inches(1.0)

            # Set column widths
            for i, width in enumerate([width_items, width_price_est, width_qty]):
                options_table.columns[i].width = width

            # Header row
            header_cells = options_table.rows[0].cells
            headers = ["Items", "Price Est", "Qty"]
            for i, cell in enumerate(header_cells):
                cell.width = [width_items, width_price_est, width_qty][i]
                cell.text = headers[i]
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell_para.runs:
                    run.bold = True

            # Add optional items rows
            for item in optional_items_data:
                row = options_table.add_row()
                cells = row.cells
                cells[0].text = item.get('description', '')
                quantity = item.get('quantity', 1)
                price_each = item.get('price_each', 0)
                is_tbd_price = item.get('is_tbd_price', False)

                # Price Est column
                if is_tbd_price:
                    cells[1].text = "TBD"
                else:
                    cells[1].text = f"${price_each:,.2f}" if price_each > 0 else "TBD"
                
                # Qty column
                cells[2].text = str(int(quantity)) if quantity == int(quantity) else str(quantity)

                # Center numeric columns
                for i in range(1, 3):
                    cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spacing after table
            doc.add_paragraph()

        # Footnote about pricing
        footnote_para = doc.add_paragraph()
        footnote_para.style = 'Caption'
        footnote_para.add_run("\nPricing valid for 60 days from the date of this proposal. Taxes not included.")


        # Add a page break to start the Proposal Details page
        doc.add_page_break()

        
        doc.add_section()
        section_index = len(doc.sections) - 1
        current_section = doc.sections[section_index]
        current_page_num = 7
        setup_document_headers_and_footers(doc, current_section, client_name, proposal_num, current_date, current_page_num)
        # Add "Proposal includes:" heading - black, underlined, smaller
        proposal_includes_para = doc.add_paragraph()
        proposal_includes_run = proposal_includes_para.add_run('Proposal includes:')
        proposal_includes_run.font.size = Pt(10)  # Smaller font size
        proposal_includes_run.font.bold = True
        proposal_includes_run.font.underline = True  # Add underline
        proposal_includes_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        # Removed empty paragraph to reduce spacing

        # Create a bulleted list of included items
        included_items = [
            "All mechanical hardware as outlined in this proposal",
            f"Freight to {{{{CUSTOMER_LOCATION}}}} (2 Calistoga Trucks, Side Load)",
            "Unloading the equipment upon arrival",
            "Moving equipment from dock area to installation area",
            "Mechanical installation using our own forklift and scissor lifts",
            "Electrical installation and controls within the equipment",
            "Configuring controls and checking the equipment for proper operation",
            "Operator and maintenance training",
            "Manuals and Documentation",
            "PIH Exclusive: 3 Scheduled Maintenance Visits over the first 2 years (Modula requires 1)",
            "2 Year Warranty Parts and Labor",
            "Seismic Anchoring Calculations and Certifications (Price to be determined)"
        ]

        # Try to extract location from customer_info or extraction
        freight_location = "Tualatin, OR"
        try:
            if hasattr(customer_info, 'get') and customer_info.get('location'):
                freight_location = customer_info.get('location')

            # Replace the placeholder with actual location if found
            included_items[1] = included_items[1].replace("{{CUSTOMER_LOCATION}}", freight_location)
        except:
            # Keep the placeholder if location can't be determined
            pass

        for item in included_items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)
            for run in p.runs:
                run.font.size = Pt(10)

        # No paragraph break to reduce spacing

        # Add "Proposal does not include:" heading - black, underlined, smaller
        proposal_excludes_para = doc.add_paragraph()
        proposal_excludes_run = proposal_excludes_para.add_run('Proposal does not include:')
        proposal_excludes_run.font.size = Pt(10)  # Smaller font size
        proposal_excludes_run.font.bold = True
        proposal_excludes_run.font.underline = True  # Add underline
        proposal_excludes_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        # No empty paragraph to reduce spacing

        # Create a bulleted list of excluded items
        # Get customer company name for excluded items
        customer_company = customer_info.get('customer_name', customer_info.get('name', 'Customer'))

        excluded_items = [
            f"Electrical Hookup to VLM provided by {customer_company}",
            "Fire Suppression",
            "Installation Permits",
            "Providing a clear path for moving equipment to the installation area",
            "Providing a clear area for installation and the erected equipment",
            "Computer Hardware and data drops",
            "Local and State Taxes on applicable items",
            f"Seismic reports may require modifications to the {customer_company} facility which will be the customer's responsibility"
        ]

        for item in excluded_items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)
            for run in p.runs:
                run.font.size = Pt(10)

        # No paragraph break to reduce spacing between excluded items and lead time

        # Add Lead Time section - more compact
        lead_time_heading = doc.add_paragraph()
        lead_time_heading.paragraph_format.space_before = Pt(0)
        lead_time_heading.paragraph_format.space_after = Pt(0)
        lead_time_heading.add_run("Lead Time").bold = True
        lead_time_heading.add_run(": 10-12 Weeks")
        for run in lead_time_heading.runs:
            run.font.size = Pt(10)

        # Add Payment Schedule section - less spacing
        payment_heading = doc.add_paragraph()
        payment_heading.paragraph_format.space_before = Pt(6)
        payment_heading.paragraph_format.space_before = Pt(0)
        payment_heading.paragraph_format.space_after = Pt(0)  # Reduce space before
        payment_heading.add_run("Payment Schedule:").bold = True
        for run in payment_heading.runs:
            run.font.size = Pt(10)

        payment_schedule = [
            "35% due upon purchase order, Net 10 days",
            "55% due upon shipping, Net 30 days",
            "10% due upon acceptance, Net 30 days"
        ]

        for item in payment_schedule:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(2.0)  # Indent for better alignment
            p.add_run(item)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.size = Pt(10)

        # No paragraph break to reduce spacing between payment schedule and notes

        # Add notes
        note1_para = doc.add_paragraph()
        note1_para.paragraph_format.space_before = Pt(0)
        note1_para.paragraph_format.space_after = Pt(0)
        note1_para.add_run("Note 1.").bold = True
        note1_para.add_run("\tAll Permits are \"By Customer\". Please note that processing time can take 3 to 5 weeks or more.")
        # Set all runs to 10pt
        for run in note1_para.runs:
            run.font.size = Pt(10)

        note2_para = doc.add_paragraph()
        note2_para.paragraph_format.space_before = Pt(0)
        note2_para.paragraph_format.space_after = Pt(0)
        note2_para.add_run("Note 2.").bold = True
        note2_para.add_run("\tSeismic calculations for the equipment are required to obtain installation permits. Seismic calculations will be prepared by an Architect selected by Pacific Integrated Handling. The purpose of the calculations is to show that the floor can support the installed storage equipment. The Architect will prepare final project calculations and drawings based on customer supplied floor data. The data needed includes floor hardness and thickness, rebar size and spacing, and the soil compaction or floor load rating and construction if not on ground floor. If the floor proves to be inadequate to support the equipment, the cost of changes to the floor and/or equipment is the responsibility of the customer.")
        # Set all runs to 10pt
        for run in note2_para.runs:
            run.font.size = Pt(10)

        # No paragraph break to reduce spacing between notes and closing

        # Add closing and signature block
        closing_para = doc.add_paragraph()
        closing_para.paragraph_format.space_before = Pt(0)
        closing_para.paragraph_format.space_after = Pt(0)
        closing_para.add_run("Best regards,")
        for run in closing_para.runs:
            run.font.size = Pt(10)

        # Add minimal space before signature block
        doc.add_paragraph()

        # Create signature block with a 2x2 table for alignment
        signature_table = doc.add_table(rows=2, cols=2)
        signature_table.autofit = False
        signature_table.style = 'Table Grid'
        signature_table.style = 'Normal Table'  # Remove borders

        # Set column widths
        signature_table.columns[0].width = Inches(3.0)
        signature_table.columns[1].width = Inches(3.5)

        # Add content to signature table
        contact_name_sig = (customer_info or {}).get('contact_name', 'Josh Jancola')
        signature_table.cell(0, 0).text = f"{contact_name_sig}, Sales Representative"
        for para in signature_table.cell(0, 0).paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

        signature_table.cell(0, 1).paragraphs[0].add_run("Accepted by: ").bold = True
        for run in signature_table.cell(0, 1).paragraphs[0].runs:
            run.font.size = Pt(10)
        signature_table.cell(0, 1).paragraphs[0].add_run("_________________________")
        for run in signature_table.cell(0, 1).paragraphs[0].runs:
            run.font.size = Pt(10)

        signature_table.cell(1, 0).text = ""

        signature_table.cell(1, 1).paragraphs[0].add_run("Signature: ").bold = True
        for run in signature_table.cell(1, 1).paragraphs[0].runs:
            run.font.size = Pt(10)
        signature_table.cell(1, 1).paragraphs[0].add_run("_________________________")
        for run in signature_table.cell(1, 1).paragraphs[0].runs:
            run.font.size = Pt(10)

        # Add date line
        date_para = doc.add_paragraph()
        date_para.paragraph_format.space_before = Pt(0)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run("Date: ").bold = True
        date_para.add_run("_____________________________")
        for run in date_para.runs:
            run.font.size = Pt(10)
        
        
        doc.add_page_break()
        
        
                        # ======== PAGE 11: SCREENSHOT FROM PDF PAGE 4 ========
        from docx.shared import Inches
        import tempfile, os, fitz

        # Add new section
        doc.add_section()
        section = doc.sections[-1]

        # Unlink & clear header/footer
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)

        try:
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            # Open PDF and crop
            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[3]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 30)
            pix = page.get_pixmap(dpi=200, clip=crop_rect)
            pix.save(temp_path)
            pdf_doc.close()

            # Later when inserting image
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")


                
            



                                # ======== PAGE 12: SCREENSHOT FROM PDF PAGE 6 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)

        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[5]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")


    
        
                                # ======== PAGE 13: SCREENSHOT FROM PDF PAGE 7 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[6]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")


                # keep margins respected
                

        
                                # ======== PAGE 14: SCREENSHOT FROM PDF PAGE 8 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[7]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
           

        
                                # ======== PAGE 15: SCREENSHOT FROM PDF PAGE 9 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[8]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
           

                
        
                                # ======== PAGE 16: SCREENSHOT FROM PDF PAGE 10 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[9]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
            

            
        
                                # ======== PAGE 17: SCREENSHOT FROM PDF PAGE 21 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
      
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[20]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
            

       
        
                                # ======== PAGE 18: SCREENSHOT FROM PDF PAGE 22 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[21]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
            

        
                                # ======== PAGE 19: SCREENSHOT FROM PDF PAGE 23 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[23]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
            
        
        
                                # ======== PAGE 20: SCREENSHOT FROM PDF PAGE 26 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[25]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
            
         
       
                                # ======== PAGE 21: SCREENSHOT FROM PDF PAGE 27 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[26]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
            
        

               
        
                                # ======== PAGE 22: SCREENSHOT FROM PDF PAGE 28 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

        # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
       
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[27]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
            
           
            
        
                                # ======== PAGE 23: SCREENSHOT FROM PDF PAGE 29 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()
             # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[28]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
            
        
        
                                # ======== PAGE 24: SCREENSHOT FROM PDF PAGE 30 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()
         # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[29]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
            
      
                                # ======== PAGE 25: SCREENSHOT FROM PDF PAGE 31 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()
         # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[30]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()

            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
       
          
        
                                # ======== PAGE 26: SCREENSHOT FROM PDF PAGE 20 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()
         # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[19]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
       
          
                                # ======== PAGE 27: SCREENSHOT FROM PDF PAGE 32 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()
         # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[31]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
          
          
        
                                # ======== PAGE 28: SCREENSHOT FROM PDF PAGE 33 ========
        doc.add_section()
        section = doc.sections[-1]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()
         # Remove all margins → no space for header/footer
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
    
        try:
            import tempfile, os, fitz
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                temp_path = temp_img.name

            pdf_doc = fitz.open(pdf_path)
            page = pdf_doc[32]  
            rect = page.rect
            crop_rect = fitz.Rect(rect.x0, rect.y0 + 80, rect.x1, rect.y1 - 50)  # crop 80px from top & bottom
            pix = page.get_pixmap(dpi=200, clip=crop_rect)

            pix.save(temp_path)
            pdf_doc.close()
            if os.path.exists(temp_path):
                usable_width = section.page_width - section.left_margin - section.right_margin
                usable_height = section.page_height - section.top_margin - section.bottom_margin

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(temp_path, width=usable_width, height=usable_height)

                os.unlink(temp_path)

            else:
                doc.add_paragraph("Image extraction failed. Please refer to the original PDF.")
        except Exception as img_e:
            doc.add_paragraph(f"Error extracting machine specifications image: {str(img_e)}")
             
  
        
        # Save the Word document
        doc.save(word_path)
        logger.debug(f"Word document created at {word_path}")

    except Exception as e:
        logger.error(f"Error processing PDF to Word: {str(e)}")

        # Create a fallback Word document with error information
        try:
            from datetime import datetime
            from docx import Document
            from docx.shared import Pt, Inches

            doc = Document()

            # Set document-wide default font to Calibri
            style = doc.styles['Normal']
            style.font.name = 'Calibri'

            # Set 1-inch margins for all sections
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            heading = doc.add_heading('PDF Extraction Error', 0)
            # Set Calibri for heading too
            for run in heading.runs:
                run.font.name = 'Calibri'

            error_para = doc.add_paragraph(f"Error processing PDF: {str(e)}")
            error_para.runs[0].font.size = Pt(11)

            file_para = doc.add_paragraph(f"File: {os.path.basename(pdf_path)}")
            file_para.runs[0].font.size = Pt(11)

            date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_para.runs[0].font.size = Pt(11)

            doc.save(word_path)
        except Exception as fallback_e:
            logger.error(f"Even fallback Word creation failed: {str(fallback_e)}")
            raise

def add_marketing_content(doc, marketing_pdf_path='static/assets/modula_vlm_marketing.pdf'):
    """
    Add marketing content from a PDF file as full-page embedded images

    Args:
        doc: The Word document to add the content to
        marketing_pdf_path: Path to the marketing PDF
    """
    logger.info(f"Adding full-page marketing content from {marketing_pdf_path}")

    try:
        if not os.path.exists(marketing_pdf_path):
            logger.error(f"Marketing PDF not found at {marketing_pdf_path}")
            return False

        # Convert the PDF to images at higher DPI for better quality
        images = convert_from_path(marketing_pdf_path, dpi=300)
        logger.info(f"Converted {len(images)} pages from marketing PDF")

        # Add each image as a full page
        for i, image in enumerate(images):
            # Add a page break before each marketing page
            page_break = doc.add_paragraph()
            run = page_break.add_run()
            run.add_break(WD_BREAK.PAGE)

            # Create a temporary file for the image
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                image_path = temp_file.name

            # Save the image
            image.save(image_path, 'PNG')

            # Create a paragraph with no margins/spacing
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Remove paragraph spacing
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = 1.0

            run = paragraph.add_run()

            # Add the image to fill the entire page (8.5 x 11 inches with minimal margins)
            from docx.shared import Inches, Pt
            run.add_picture(image_path, width=Inches(8.0), height=Inches(10.5))

            # Clean up the temporary file
            os.unlink(image_path)

        logger.info(f"Successfully added {len(images)} full-page marketing pages to document")
        return True
    except Exception as e:
        logger.error(f"Error adding marketing content: {str(e)}")
        return False
"""
This script integrates the Modula VLM marketing PDF pages directly into
the generated Word document as embedded images.
"""
import os
import logging
import tempfile
from pdf2image import convert_from_path

import shutil

# Setup logging s
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def modify_pdf_processor():
    """
    Modify the PDF processor to include marketing PDF pages as images
    in the generated Word document.
    """
    pdf_processor_path = "utils/pdf_processor.py"
    
    if not os.path.exists(pdf_processor_path):
        logger.error(f"Could not find {pdf_processor_path}")
        return False
    
    logger.info(f"Reading {pdf_processor_path}...")
    
    with open(pdf_processor_path, 'r') as f:
        content = f.read()
    
    # Find the end of the imports section
    import_section_end = content.find("# Set up logging")
    if import_section_end == -1:
        import_section_end = content.find("logger = logging.getLogger")
    
    # Add needed imports if they don't exist
    imports_to_add = ""
    if "from pdf2image import convert_from_path" not in content:
        imports_to_add += "from pdf2image import convert_from_path\n"
    if "from docx.enum.section import WD_SECTION" not in content:
        imports_to_add += "from docx.enum.section import WD_SECTION\n"
    if "import tempfile" not in content:
        imports_to_add += "import tempfile\n"
    
    if imports_to_add:
        content = content[:import_section_end] + imports_to_add + content[import_section_end:]
        logger.info("Added necessary imports")
    
    # Find the end of the process_pdf_to_word function
    word_func_start = content.find("def process_pdf_to_word(")
    if word_func_start == -1:
        logger.error("Could not find process_pdf_to_word function")
        return False
    
    # Find where the document is saved - this is where we will insert our code
    save_line = content.find("doc.save(word_path)", word_func_start)
    if save_line == -1:
        logger.error("Could not find where document is saved")
        return False
    
    # See if our code is already there
    if "add_marketing_content" in content[word_func_start:save_line]:
        logger.info("Marketing content code already present")
    else:
        # Add the code to embed marketing PDF before saving
        insert_code = """
        # Add marketing content from PDF
        add_marketing_content(doc)
        """
        
        # Insert the code just before the save line
        # First, find the line start by going back to the previous newline
        line_start = content.rfind("\\n", 0, save_line) + 1
        
        # Extract leading whitespace from the save line for proper indentation
        indentation = ""
        for char in content[line_start:save_line]:
            if char == ' ' or char == '\\t':
                indentation += char
            else:
                break
        
        # Format our insert code with the same indentation
        formatted_insert = insert_code.replace("        ", indentation)
        
        # Insert the code
        content = content[:line_start] + formatted_insert + content[line_start:]
        logger.info("Added code to embed marketing PDF before saving")
    
    # Add the function to embed marketing content if it doesn't exist
    if "def add_marketing_content(" not in content:
        marketing_func = """
def add_marketing_content(doc, marketing_pdf_path='static/assets/modula_vlm_marketing.pdf'):
    \"\"\"
    Add marketing content from a PDF file as embedded images
    
    Args:
        doc: The Word document to add the content to
        marketing_pdf_path: Path to the marketing PDF
    \"\"\"
    logger.info(f"Adding marketing content from {marketing_pdf_path}")
    
    try:
        if not os.path.exists(marketing_pdf_path):
            logger.error(f"Marketing PDF not found at {marketing_pdf_path}")
            return False
        
        # Convert the PDF to images
        images = convert_from_path(marketing_pdf_path, dpi=200)
        logger.info(f"Converted {len(images)} pages from marketing PDF")
        
        # Add a section break before the marketing content
        last_paragraph = doc.add_paragraph()
        run = last_paragraph.add_run()
        run.add_break(WD_SECTION.NEW_PAGE)
        
        # Add a title for the marketing section
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        title = doc.add_heading("Modula VLM Additional Information", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add each image on a new page
        for i, image in enumerate(images):
            # Add a page break after the first marketing page
            if i > 0:
                page_break = doc.add_paragraph()
                run = page_break.add_run()
                run.add_break(WD_SECTION.NEW_PAGE)
            
            # Create a temporary file for the image
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                image_path = temp_file.name
                
            # Save the image
            image.save(image_path, 'PNG')
            
            # Add the image
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            
            # Add the image with proper width (7 inches)
            from docx.shared import Inches
            run.add_picture(image_path, width=Inches(7))
            
            # Clean up the temporary file
            os.unlink(image_path)
        
        logger.info(f"Successfully added {len(images)} marketing pages to document")
        return True
    except Exception as e:
        logger.error(f"Error adding marketing content: {str(e)}")
        return False
"""
        
        # Add the function to the end of the file
        content += "\n" + marketing_func
        logger.info("Added add_marketing_content function")
    
    # Write the modified content back to the file
    with open(pdf_processor_path, 'w') as f:
        f.write(content)
    
    logger.info("Successfully modified PDF processor")
    return True

def clean_up_pdf_downloads():
    """
    Remove PDF download options and update the UI to indicate 
    marketing content is included in the Word document.
    """
    # Find the result.html template
    template_path = "templates/result.html"
    
    if not os.path.exists(template_path):
        logger.error(f"Could not find {template_path}")
        return False
    
    logger.info(f"Reading {template_path}...")
    
    with open(template_path, 'r') as f:
        content = f.read()
    
    # Check if we already modified the template
    if "PDF with Marketing" in content:
        # Remove PDF download button
        import re
        pdf_button_pattern = r'<a href="javascript:void\(0\)" id="pdf-download-btn".*?Download PDF with Marketing.*?</a>'
        content = re.sub(pdf_button_pattern, "", content, flags=re.DOTALL)
        
        # Update Word button text
        content = content.replace(
            'Download Word Document',
            'Download Document (with Marketing Pages)'
        )
        
        # Update modal title and messages
        if 'Download Now' in content:
            content = content.replace(
                'Download Now',
                'Download Document'
            )
        
        # Also update the modal message to indicate marketing pages are included
        modal_body_pattern = r'<div class="modal-body">(.*?)</div>'
        modal_body_match = re.search(modal_body_pattern, content, re.DOTALL)
        
        if modal_body_match:
            old_modal_body = modal_body_match.group(0)
            new_modal_body = old_modal_body.replace(
                "</p>",
                "</p>\n                <p>This document includes the complete marketing package as additional pages.</p>"
            )
            content = content.replace(old_modal_body, new_modal_body)
        
        # Write the modified content back to the file
        with open(template_path, 'w') as f:
            f.write(content)
        
        logger.info("Updated UI to remove PDF download option")
        return True
    
    # Just in case the PDF button wasn't added yet, update the Word button text
    if 'Download Document' in content:
        content = content.replace(
            'Download Document',
            'Download Document (with Marketing Pages)'
        )
        
        # Write the modified content back to the file
        with open(template_path, 'w') as f:
            f.write(content)
        
        logger.info("Updated Word button text")
    
    return True

if __name__ == "__main__":
    # Make sure marketing PDF exists
    marketing_pdf_path = "static/assets/modula_vlm_marketing.pdf"
    if not os.path.exists(marketing_pdf_path):
        logger.error(f"Marketing PDF not found at {marketing_pdf_path}")
        print(f"Error: Marketing PDF not found at {marketing_pdf_path}")
        print("Please ensure the marketing PDF is located at this path.")
        exit(1)
    
    # Modify the PDF processor
    if modify_pdf_processor():
        # Remove PDF download options
        clean_up_pdf_downloads()
        
        print("Successfully implemented marketing PDF integration!")
        print("Changes made:")
        print("1. Modified PDF processor to embed marketing PDF pages as images")
        print("2. Removed separate PDF download option")
        print("3. Updated UI to indicate marketing pages are included in the Word document")
        print("\nThe Word document now includes all marketing content as embedded images.")
    else:
        print("Failed to modify PDF processor. Please check the logs for details.")
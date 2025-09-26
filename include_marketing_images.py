"""
This script adds the Modula VLM marketing PDF as embedded images
directly into the generated Word document.

Instead of appending the PDF, this approach converts each page to a PNG
and adds them to the end of the Word document.
"""
import os
import logging
import tempfile
import shutil
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def convert_pdf_to_images(pdf_path, output_dir=None, dpi=200):
    """
    Convert each page of a PDF to a PNG image
    
    Args:
        pdf_path: Path to the PDF file
        output_dir: Directory to save the images to (or a temp dir is created if None)
        dpi: Resolution of the output images
        
    Returns:
        List of paths to the generated image files
    """
    logger.info(f"Converting PDF {pdf_path} to images with DPI {dpi}")
    
    if not os.path.exists(pdf_path):
        logger.error(f"PDF file not found: {pdf_path}")
        return []
    
    # Create a temporary directory if output_dir is not provided
    if output_dir is None:
        temp_dir = tempfile.mkdtemp()
        output_dir = temp_dir
    else:
        temp_dir = None
        os.makedirs(output_dir, exist_ok=True)
    
    image_paths = []
    
    try:
        # Convert PDF to images
        images = convert_from_path(pdf_path, dpi=dpi)
        logger.info(f"Converted {len(images)} pages from PDF")
        
        # Save each image
        for i, image in enumerate(images):
            image_path = os.path.join(output_dir, f"page_{i+1}.png")
            image.save(image_path, "PNG")
            image_paths.append(image_path)
            logger.info(f"Saved page {i+1} to {image_path}")
        
        return image_paths
    except Exception as e:
        logger.error(f"Error converting PDF to images: {str(e)}")
        return []
    finally:
        # Clean up temporary directory if created
        if temp_dir and not image_paths:
            try:
                shutil.rmtree(temp_dir)
                logger.info(f"Removed temporary directory {temp_dir}")
            except Exception as e:
                logger.error(f"Error removing temporary directory: {str(e)}")

def add_images_to_word(docx_path, image_paths, marketing_title="Modula VLM Additional Information"):
    """
    Add images to an existing Word document
    
    Args:
        docx_path: Path to the Word document
        image_paths: List of paths to the images to add
        marketing_title: Title to add before the marketing images
    """
    logger.info(f"Adding {len(image_paths)} images to {docx_path}")
    
    try:
        # Open the document
        doc = Document(docx_path)
        
        # Add a section break before the marketing content
        last_paragraph = doc.add_paragraph()
        run = last_paragraph.add_run()
        run.add_break(WD_SECTION.NEW_PAGE)
        
        # Add a title for the marketing section
        title = doc.add_heading(marketing_title, level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add each image on a new page
        for i, image_path in enumerate(image_paths):
            # Add a page break after the first marketing page
            if i > 0:
                page_break = doc.add_paragraph()
                run = page_break.add_run()
                run.add_break(WD_SECTION.NEW_PAGE)
            
            # Add the image
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            
            # Use 7 inches for image width to fit on the page with margins
            run.add_picture(image_path, width=Inches(7))
            logger.info(f"Added image {i+1} to document")
        
        # Save the document
        doc.save(docx_path)
        logger.info(f"Saved document with {len(image_paths)} marketing images")
        
        return True
    except Exception as e:
        logger.error(f"Error adding images to Word document: {str(e)}")
        return False

def modify_pdf_processor():
    """Modify the PDF processor to include marketing PDF pages as images"""
    pdf_processor_file = "utils/pdf_processor.py"
    
    # Check if the file exists
    if not os.path.exists(pdf_processor_file):
        logger.error(f"PDF processor file not found: {pdf_processor_file}")
        return False
    
    # Read the file
    with open(pdf_processor_file, "r") as f:
        content = f.read()
    
    # Add necessary imports if they don't exist
    add_imports_to_processor(pdf_processor_file)
    
    # Check if marketing code is already added
    if "def add_marketing_images_to_document" in content:
        logger.info("Marketing images code already exists in PDF processor")
        return True
    
    # Find where to add the function
    function_to_add = """
def add_marketing_images_to_document(doc, marketing_pdf_path='static/assets/modula_vlm_marketing.pdf', dpi=200):
    \"\"\"
    Add marketing images from PDF to Word document
    
    Args:
        doc: Word document object
        marketing_pdf_path: Path to marketing PDF
        dpi: Image resolution
    
    Returns:
        bool: Success or failure
    \"\"\"
    logger.info(f"Adding marketing images from {marketing_pdf_path}")
    
    try:
        if not os.path.exists(marketing_pdf_path):
            logger.error(f"Marketing PDF not found: {marketing_pdf_path}")
            return False
        
        # Create a temp directory for images
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Convert PDF to images
            images = convert_from_path(marketing_pdf_path, dpi=dpi)
            logger.info(f"Converted {len(images)} pages from marketing PDF")
            
            # Add a section break before the marketing content
            last_paragraph = doc.add_paragraph()
            run = last_paragraph.add_run()
            run.add_break(WD_SECTION.NEW_PAGE)
            
            # Add a title for the marketing section
            title = doc.add_heading("Modula VLM Additional Information", level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add each image on a new page
            for i, image in enumerate(images):
                # Add a page break after the first marketing page
                if i > 0:
                    page_break = doc.add_paragraph()
                    run = page_break.add_run()
                    run.add_break(WD_SECTION.NEW_PAGE)
                
                # Create a temporary file for the image
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                    image_path = temp_file.name
                
                # Save the image
                image.save(image_path, 'PNG')
                
                # Add the image
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()
                
                # Add the image with proper width (7 inches)
                run.add_picture(image_path, width=Inches(7))
                
                # Clean up the temporary file
                os.unlink(image_path)
            
            logger.info(f"Successfully added {len(images)} marketing pages to document")
            return True
        finally:
            # Clean up the temp directory
            shutil.rmtree(temp_dir)
            logger.info(f"Removed temporary directory {temp_dir}")
    
    except Exception as e:
        logger.error(f"Error adding marketing images: {str(e)}")
        return False
"""
    
    # Add the function to the end of the file
    with open(pdf_processor_file, "a") as f:
        f.write(function_to_add)
    
    logger.info("Added marketing images function to PDF processor")
    
    # Now modify the process_pdf_to_word function to call our new function
    with open(pdf_processor_file, "r") as f:
        content = f.read()
    
    # Find the save line in process_pdf_to_word
    save_line_index = content.find("doc.save(word_path)")
    if save_line_index == -1:
        logger.error("Could not find doc.save line in process_pdf_to_word function")
        return False
    
    # Find the start of that line
    line_start = content.rfind("\n", 0, save_line_index) + 1
    
    # Get the indentation
    indentation = ""
    for char in content[line_start:save_line_index]:
        if char in [" ", "\t"]:
            indentation += char
        else:
            break
    
    # Check if the marketing code is already added
    marketing_call = indentation + "# Add marketing content\n"
    marketing_call += indentation + "add_marketing_images_to_document(doc)\n"
    
    if "add_marketing_images_to_document(doc)" in content:
        logger.info("Marketing images call already exists in process_pdf_to_word function")
    else:
        # Insert our code before the save line
        new_content = content[:line_start] + marketing_call + content[line_start:]
        
        # Write the modified content back to the file
        with open(pdf_processor_file, "w") as f:
            f.write(new_content)
        
        logger.info("Modified process_pdf_to_word function to add marketing images")
    
    return True

def add_imports_to_processor(pdf_processor_file):
    """Add necessary imports to the PDF processor file"""
    with open(pdf_processor_file, "r") as f:
        content = f.read()
    
    # Prepare the imports to add
    imports_to_add = []
    
    if "from pdf2image import convert_from_path" not in content:
        imports_to_add.append("from pdf2image import convert_from_path")
    
    if "import tempfile" not in content:
        imports_to_add.append("import tempfile")
    
    if "import shutil" not in content:
        imports_to_add.append("import shutil")
    
    if "from docx.enum.section import WD_SECTION" not in content:
        imports_to_add.append("from docx.enum.section import WD_SECTION")
    
    # Find where to add the imports
    if imports_to_add:
        # Find the end of the import section
        import_section_end = content.find("# Set up logging")
        if import_section_end == -1:
            import_section_end = content.find("logger = logging.getLogger")
        
        if import_section_end != -1:
            # Add the imports
            imports_text = "\n".join(imports_to_add) + "\n"
            
            new_content = content[:import_section_end] + imports_text + content[import_section_end:]
            
            # Write the modified content back to the file
            with open(pdf_processor_file, "w") as f:
                f.write(new_content)
            
            logger.info(f"Added imports to PDF processor: {imports_to_add}")
        else:
            logger.error("Could not find where to add imports in PDF processor")
    else:
        logger.info("All required imports already exist in PDF processor")

def clean_up_pdf_download_option():
    """Remove PDF download option from the result page"""
    template_file = "templates/result.html"
    
    if not os.path.exists(template_file):
        logger.error(f"Template file not found: {template_file}")
        return False
    
    with open(template_file, "r") as f:
        content = f.read()
    
    # Update the Word button text to indicate it includes marketing materials
    old_word_button = '<button id="word-download-btn" class="btn btn-primary">Download Word Document</button>'
    new_word_button = '<button id="word-download-btn" class="btn btn-primary">Download Document (includes Marketing)</button>'
    
    # Check if the PDF button exists
    pdf_button_start = content.find('id="pdf-download-btn"')
    
    if pdf_button_start != -1:
        # Find the start and end of the button element
        button_start = content.rfind("<", 0, pdf_button_start)
        button_end = content.find(">", pdf_button_start) + 1
        
        # Remove the PDF button
        content = content.replace(content[button_start:button_end], "")
        
        # Replace any remaining references to the PDF download
        content = content.replace(".pdf", "")
        content = content.replace("PDF", "")
        
        # Update the Word button text
        content = content.replace(old_word_button, new_word_button)
        
        # Write the modified content back
        with open(template_file, "w") as f:
            f.write(content)
        
        logger.info("Removed PDF download option from result page")
    else:
        # Just update the Word button text
        if old_word_button in content:
            content = content.replace(old_word_button, new_word_button)
            
            with open(template_file, "w") as f:
                f.write(content)
            
            logger.info("Updated Word button text")
        else:
            logger.info("Word button text already updated")
    
    return True

if __name__ == "__main__":
    # Check if the marketing PDF exists
    marketing_pdf_path = "static/assets/modula_vlm_marketing.pdf"
    
    if not os.path.exists("static/assets"):
        os.makedirs("static/assets", exist_ok=True)
        logger.info("Created static/assets directory")
    
    if not os.path.exists(marketing_pdf_path):
        logger.warning(f"Marketing PDF not found at {marketing_pdf_path}")
        logger.warning("Please copy the marketing PDF to this location")
        print(f"Marketing PDF not found at {marketing_pdf_path}")
        print("Please copy the marketing PDF to this location")
    
    # Modify the PDF processor
    if modify_pdf_processor():
        print("Successfully modified PDF processor to include marketing images")
    else:
        print("Failed to modify PDF processor")
    
    # Clean up PDF download option
    if clean_up_pdf_download_option():
        print("Successfully removed PDF download option")
    else:
        print("Failed to remove PDF download option")
    
    print("\nImplementation complete. The Word document will now include marketing materials as embedded images.")